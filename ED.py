# =========================
# THƯ VIỆN BẮT BUỘC VÀ BỔ SUNG
# =========================
from datetime import datetime
import os
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier, StackingClassifier
from sklearn.metrics import (
    confusion_matrix,
    f1_score,
    accuracy_score,
    recall_score,
    precision_score,
    roc_auc_score,
    ConfusionMatrixDisplay,
)
from xgboost import XGBClassifier
import time

# Thư viện RSS Feed
try:
    import feedparser
    _FEEDPARSER_OK = True
except Exception:
    feedparser = None
    _FEEDPARSER_OK = False

# Thư viện GOOGLE GEMINI VÀ OPENAI (Giữ nguyên logic kiểm tra thư viện)
try:
    from google import genai
    from google.genai.errors import APIError
    _GEMINI_OK = True
except Exception:
    genai = None
    APIError = Exception
    _GEMINI_OK = False

try:
    from openai import OpenAI
    _OPENAI_OK = True
except Exception:
    OpenAI = None
    _OPENAI_OK = False

# Thư viện Word Export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    _WORD_OK = True
except Exception:
    _WORD_OK = False

MODEL_NAME = "gemini-2.5-flash"

# =========================
# HÀM PHÂN LOẠI PD THEO 5 CẤP ĐỘ
# =========================

def classify_pd(pd_value):
    """
    Phân loại PD theo 5 cấp độ với rating và màu sắc gradient.

    Args:
        pd_value: Xác suất vỡ nợ (0-1)

    Returns:
        dict: {
            'range': 'PD Range',
            'classification': 'Phân loại',
            'rating': 'Rating (AAA-D)',
            'meaning': 'Ý nghĩa',
            'color': 'Mã màu hex',
            'gradient_color': 'Gradient color'
        }
    """
    if pd.isna(pd_value):
        return {
            'range': 'N/A',
            'classification': 'Không xác định',
            'rating': 'N/A',
            'meaning': 'Thiếu dữ liệu',
            'color': '#6c757d',
            'gradient_color': 'linear-gradient(135deg, #6c757d 0%, #95a5a6 100%)'
        }

    pd_percent = pd_value * 100  # Convert to percentage

    if pd_percent < 2:
        return {
            'range': '< 2%',
            'classification': 'Rất thấp',
            'rating': 'AAA-AA',
            'meaning': 'Doanh nghiệp xuất sắc',
            'color': '#28a745',  # Green
            'gradient_color': 'linear-gradient(135deg, #28a745 0%, #20c997 100%)'
        }
    elif pd_percent < 5:
        return {
            'range': '2-5%',
            'classification': 'Thấp',
            'rating': 'A-BBB',
            'meaning': 'Doanh nghiệp tốt',
            'color': '#5cb85c',  # Light green
            'gradient_color': 'linear-gradient(135deg, #5cb85c 0%, #4cae4c 100%)'
        }
    elif pd_percent < 10:
        return {
            'range': '5-10%',
            'classification': 'Trung bình',
            'rating': 'BB',
            'meaning': 'Cần theo dõi',
            'color': '#ffc107',  # Yellow/Warning
            'gradient_color': 'linear-gradient(135deg, #ffc107 0%, #ffca2c 100%)'
        }
    elif pd_percent < 20:
        return {
            'range': '10-20%',
            'classification': 'Cao',
            'rating': 'B',
            'meaning': 'Rủi ro đáng kể',
            'color': '#fd7e14',  # Orange
            'gradient_color': 'linear-gradient(135deg, #fd7e14 0%, #ff851b 100%)'
        }
    else:  # >= 20%
        return {
            'range': '> 20%',
            'classification': 'Rất cao',
            'rating': 'CCC-D',
            'meaning': 'Nguy cơ vỡ nợ cao',
            'color': '#dc3545',  # Red
            'gradient_color': 'linear-gradient(135deg, #dc3545 0%, #c82333 100%)'
        }

# =========================
# HÀM TẠO WORD REPORT
# =========================

def generate_word_report(ratios_display, pd_value, pd_label, ai_analysis, fig_bar, fig_radar, company_name="KHÁCH HÀNG DOANH NGHIỆP"):
    """
    Tạo báo cáo Word chuyên nghiệp từ kết quả phân tích tín dụng.

    Parameters:
    - ratios_display: DataFrame chứa 14 chỉ số tài chính (index = tên chỉ số, column = giá trị)
    - pd_value: Xác suất vỡ nợ (PD) dưới dạng số float (0-1) hoặc NaN
    - pd_label: Nhãn dự đoán ("Default" hoặc "Non-Default")
    - ai_analysis: Text phân tích từ AI
    - fig_bar: Matplotlib figure của bar chart
    - fig_radar: Matplotlib figure của radar chart
    - company_name: Tên công ty (mặc định)

    Returns:
    - BytesIO object chứa Word document
    """

    if not _WORD_OK:
        raise Exception("Thiếu thư viện python-docx. Vui lòng cài đặt: pip install python-docx Pillow")

    # Tạo document mới
    doc = Document()

    # Cấu hình margin cho document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== 1. HEADER VỚI LOGO VÀ TIÊU ĐỀ =====
    # Thêm logo nếu có
    try:
        if os.path.exists("logo-agribank.jpg"):
            doc.add_picture("logo-agribank.jpg", width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    # Tiêu đề chính
    title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ RỦI RO TÍN DỤNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(194, 24, 91)  # #c2185b
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph('Dự báo Xác suất Vỡ nợ KHDN (PD) & Phân tích AI Chuyên sâu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d
    subtitle_run.font.bold = True

    # Thông tin thời gian
    date_info = doc.add_paragraph(f"Ngày xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_info.runs[0]
    date_run.font.size = Pt(10)

    # Thông tin khách hàng
    company_info = doc.add_paragraph()
    company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_info.add_run(f"Tên khách hàng: {company_name}")
    company_run.font.size = Pt(11)
    company_run.font.bold = True

    doc.add_paragraph()  # Spacer

    # ===== 2. KẾT QUẢ DỰ BÁO PD =====
    heading1 = doc.add_heading('1. KẾT QUẢ DỰ BÁO XÁC SUẤT VỠ NỢ (PD)', level=1)
    heading1_run = heading1.runs[0]
    heading1_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    pd_para = doc.add_paragraph()
    if pd.notna(pd_value):
        pd_para.add_run(f"Xác suất Vỡ nợ (PD): ").bold = True
        pd_para.add_run(f"{pd_value:.2%}\n")
        pd_para.add_run("Phân loại: ").bold = True
        pd_para.add_run(f"{pd_label}\n")

        if "Default" in pd_label and "Non-Default" not in pd_label:
            risk_run = pd_para.add_run("⚠️ RỦI RO CAO - CẦN XEM XÉT KỸ LƯỠNG")
            risk_run.bold = True
            risk_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
        else:
            safe_run = pd_para.add_run("✓ RỦI RO THẤP - KHẢ QUAN")
            safe_run.bold = True
            safe_run.font.color.rgb = RGBColor(40, 167, 69)  # Green
    else:
        pd_para.add_run("Xác suất Vỡ nợ (PD): ").bold = True
        pd_para.add_run("Không có dữ liệu")

    doc.add_paragraph()  # Spacer

    # ===== 3. BẢNG CHỈ SỐ TÀI CHÍNH =====
    heading2 = doc.add_heading('2. CHỈ SỐ TÀI CHÍNH CHI TIẾT', level=1)
    heading2_run = heading2.runs[0]
    heading2_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    # Tạo bảng
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ số Tài chính'
    hdr_cells[1].text = 'Giá trị'

    # Style header
    for cell in hdr_cells:
        cell_para = cell.paragraphs[0]
        cell_run = cell_para.runs[0]
        cell_run.font.bold = True
        cell_run.font.size = Pt(11)
        cell_run.font.color.rgb = RGBColor(255, 255, 255)
        # Set background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FF6B9D')  # Pink
        cell._element.get_or_add_tcPr().append(shading_elm)
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, row in ratios_display.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        value = row['Giá trị']
        row_cells[1].text = f"{value:.4f}" if pd.notna(value) else "N/A"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()  # Spacer

    # ===== 4. BIỂU ĐỒ VISUALIZATION =====
    doc.add_page_break()
    heading3 = doc.add_heading('3. TRỰC QUAN HÓA DỮ LIỆU', level=1)
    heading3_run = heading3.runs[0]
    heading3_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    # Bar chart
    try:
        doc.add_heading('3.1. Biểu đồ Cột - Giá trị các Chỉ số', level=2)
        bar_buffer = BytesIO()
        fig_bar.savefig(bar_buffer, format='png', dpi=150, bbox_inches='tight')
        bar_buffer.seek(0)
        doc.add_picture(bar_buffer, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Spacer
    except Exception as e:
        doc.add_paragraph(f"Không thể tạo biểu đồ cột: {str(e)}")

    # Radar chart
    try:
        doc.add_heading('3.2. Biểu đồ Radar - Phân tích Đa chiều', level=2)
        radar_buffer = BytesIO()
        fig_radar.savefig(radar_buffer, format='png', dpi=150, bbox_inches='tight')
        radar_buffer.seek(0)
        doc.add_picture(radar_buffer, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"Không thể tạo biểu đồ radar: {str(e)}")

    # ===== 5. PHÂN TÍCH AI =====
    doc.add_page_break()
    heading4 = doc.add_heading('4. PHÂN TÍCH AI & KHUYẾN NGHỊ TÍN DỤNG', level=1)
    heading4_run = heading4.runs[0]
    heading4_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    if ai_analysis and ai_analysis.strip():
        # Chia thành các đoạn và thêm vào document
        analysis_paragraphs = ai_analysis.split('\n')
        for para_text in analysis_paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                # Highlight keywords
                if "CHO VAY" in para_text and "KHÔNG CHO VAY" not in para_text:
                    for run in para.runs:
                        if "CHO VAY" in run.text:
                            run.font.color.rgb = RGBColor(40, 167, 69)  # Green
                            run.bold = True
                elif "KHÔNG CHO VAY" in para_text:
                    for run in para.runs:
                        if "KHÔNG CHO VAY" in run.text:
                            run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                            run.bold = True
    else:
        doc.add_paragraph("Chưa có phân tích từ AI. Vui lòng click nút 'Yêu cầu AI Phân tích & Đề xuất' để nhận khuyến nghị.")

    # ===== 6. FOOTER =====
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Báo cáo này được tạo tự động bởi Hệ thống Đánh giá Rủi ro Tín dụng - Powered by AI & Machine Learning\n"
        f"© {datetime.now().year} Credit Risk Assessment System | Version 2.0 Premium"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Grey

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# CẤU HÌNH TRANG (NÂNG CẤP GIAO DIỆN)
# =========================
st.set_page_config(
    page_title="Credit Risk PD & Gemini Analysis | Banking Suite",
    page_icon="🏦",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========================================
# CSS NÂNG CẤP - PHONG CÁCH NGÂN HÀNG HIỆN ĐẠI
# ========================================
st.markdown("""
<style>
/* ========== IMPORT GOOGLE FONTS ========== */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700;900&family=Playfair+Display:wght@700;900&display=swap');

/* ========== GENERAL SETTINGS ========== */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}

* {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
}

/* Main content area */
.main {
    background: linear-gradient(135deg, #fff5f7 0%, #ffe8f0 100%);
    animation: fadeIn 0.8s ease-in;
}

@keyframes fadeIn {
    from { opacity: 0; transform: translateY(20px); }
    to { opacity: 1; transform: translateY(0); }
}

/* ========== SIDEBAR TOGGLE BUTTON - LÀM NỔI BẬT ========== */
/* Làm nổi bật nút ẩn/hiện sidebar mặc định của Streamlit */
button[kind="header"] {
    background: rgba(255, 255, 255, 0.2) !important;
    border-radius: 8px !important;
    padding: 8px !important;
    transition: all 0.3s ease !important;
}

button[kind="header"]:hover {
    background: rgba(255, 255, 255, 0.3) !important;
    transform: scale(1.1) !important;
}

/* Thêm tooltip cho nút sidebar */
button[kind="header"]::after {
    content: '';
    position: absolute;
    pointer-events: none;
}

/* Style cho nút collapse khi sidebar đang mở */
[data-testid="stSidebar"][aria-expanded="true"] + div button[kind="header"] {
    box-shadow: 0 0 15px rgba(255, 255, 255, 0.5) !important;
}

/* ========== PREMIUM HEADER BANNER ========== */
.banner-title-container {
    background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 50%, #ff6b9d 100%);
    padding: 40px 50px;
    border-radius: 20px;
    box-shadow: 0 10px 40px rgba(255, 107, 157, 0.3),
                0 5px 15px rgba(255, 133, 161, 0.2),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
    margin-bottom: 30px;
    text-align: center;
    position: relative;
    overflow: hidden;
}

/* Shine effect */
.banner-title-container::before {
    content: '';
    position: absolute;
    top: -50%;
    left: -50%;
    width: 200%;
    height: 200%;
    background: linear-gradient(
        45deg,
        transparent 30%,
        rgba(255, 255, 255, 0.1) 50%,
        transparent 70%
    );
    animation: shine 3s infinite;
}

@keyframes shine {
    0% { transform: translateX(-100%) translateY(-100%) rotate(45deg); }
    100% { transform: translateX(100%) translateY(100%) rotate(45deg); }
}

.banner-title-container h1 {
    color: #ffffff !important;
    font-family: 'Playfair Display', serif !important;
    font-weight: 900 !important;
    font-size: 2.8rem !important;
    text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.3),
                 0 0 30px rgba(255, 182, 193, 0.5);
    margin-bottom: 10px !important;
    letter-spacing: -0.5px;
    position: relative;
    z-index: 1;
    animation: titleGlow 2s ease-in-out infinite alternate;
}

@keyframes titleGlow {
    from { text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.3), 0 0 30px rgba(255, 182, 193, 0.5); }
    to { text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.3), 0 0 40px rgba(255, 182, 193, 0.7); }
}

.banner-title-container h3 {
    color: #fff0f5 !important;
    font-weight: 600 !important;
    font-size: 1.3rem !important;
    margin-top: 0 !important;
    border-bottom: none !important;
    letter-spacing: 1px;
    text-transform: uppercase;
    position: relative;
    z-index: 1;
}

/* Gold accent line */
.banner-title-container::after {
    content: '';
    position: absolute;
    bottom: 0;
    left: 50%;
    transform: translateX(-50%);
    width: 80%;
    height: 3px;
    background: linear-gradient(90deg, transparent, #ffb3c6, transparent);
    z-index: 1;
}

/* ========== SIDEBAR PREMIUM STYLING ========== */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #ff6b9d 0%, #e91e63 100%) !important;
    box-shadow: 2px 0 20px rgba(0, 0, 0, 0.1);
}

[data-testid="stSidebar"] * {
    color: #ffffff !important;
}

[data-testid="stSidebar"] .stMarkdown {
    color: #e8f4f8 !important;
}

/* File uploader trong sidebar */
div[data-testid="stFileUploader"] {
    background: rgba(255, 255, 255, 0.05);
    border: 2px dashed #ffb3c6 !important;
    border-radius: 15px;
    padding: 20px;
    margin-top: 15px;
    transition: all 0.3s ease;
    backdrop-filter: blur(10px);
}

div[data-testid="stFileUploader"]:hover {
    background: rgba(255, 255, 255, 0.1);
    border-color: #ffc0cb !important;
    transform: translateY(-2px);
    box-shadow: 0 5px 20px rgba(255, 179, 198, 0.3);
}

/* ========== TABS PREMIUM DESIGN ========== */
button[data-testid="stTab"] {
    background: linear-gradient(135deg, #ffffff 0%, #fff5f7 100%);
    border: 2px solid #ffd4dd;
    border-radius: 12px 12px 0 0 !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    font-weight: 700;
    font-size: 1rem;
    color: #4a5568;
    padding: 15px 30px;
    margin-right: 8px;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
}

button[data-testid="stTab"]:hover {
    background: linear-gradient(135deg, #ffe8f0 0%, #ffd4dd 100%);
    color: #c2185b;
    border-color: #ff6b9d;
    transform: translateY(-3px);
    box-shadow: 0 5px 15px rgba(255, 107, 157, 0.2);
}

button[data-testid="stTab"][aria-selected="true"] {
    background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 100%) !important;
    color: #ffffff !important;
    border-color: #ffb3c6 !important;
    border-bottom: 3px solid #ffb3c6 !important;
    box-shadow: 0 8px 20px rgba(255, 107, 157, 0.4),
                inset 0 1px 0 rgba(255, 255, 255, 0.2);
    transform: translateY(-3px);
}

/* ========== HEADINGS ========== */
h1, h2, h3, h4 {
    color: #1a2332 !important;
    font-weight: 700 !important;
}

h2 {
    color: #c2185b !important;
    border-bottom: 3px solid #ffb3c6;
    padding-bottom: 10px;
    margin-bottom: 20px !important;
}

h3 {
    color: #ff6b9d !important;
    border-bottom: 2px solid rgba(255, 179, 198, 0.3);
    padding-bottom: 8px;
    margin-bottom: 15px !important;
}

/* ========== METRIC CONTAINERS ========== */
div[data-testid="metric-container"] {
    background: linear-gradient(135deg, #ffffff 0%, #fff5f7 100%);
    border: 2px solid transparent;
    border-image: linear-gradient(135deg, #ffb3c6, #ff6b9d) 1;
    border-radius: 16px;
    padding: 20px;
    box-shadow: 0 8px 25px rgba(255, 107, 157, 0.12),
                0 3px 10px rgba(0, 0, 0, 0.08);
    transition: all 0.3s ease;
}

div[data-testid="metric-container"]:hover {
    transform: translateY(-5px) scale(1.02);
    box-shadow: 0 12px 35px rgba(255, 107, 157, 0.2),
                0 5px 15px rgba(255, 179, 198, 0.15);
}

/* Metric label */
div[data-testid="metric-container"] label {
    font-weight: 700 !important;
    color: #c2185b !important;
    font-size: 0.9rem !important;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

/* Metric value */
div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: #ff6b9d !important;
    font-weight: 900 !important;
    font-size: 2.2rem !important;
}

/* ========== BUTTONS PREMIUM ========== */
button[kind="primary"] {
    background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 100%) !important;
    border: 2px solid #ffb3c6 !important;
    border-radius: 12px !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    font-size: 1.05rem !important;
    padding: 12px 30px !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    box-shadow: 0 6px 20px rgba(255, 107, 157, 0.3),
                inset 0 1px 0 rgba(255, 255, 255, 0.2);
    text-transform: uppercase;
    letter-spacing: 1px;
}

button[kind="primary"]:hover {
    background: linear-gradient(135deg, #e91e63 0%, #f06292 100%) !important;
    border-color: #ffc0cb !important;
    transform: translateY(-2px) scale(1.02);
    box-shadow: 0 10px 30px rgba(255, 107, 157, 0.4),
                0 5px 15px rgba(255, 179, 198, 0.3),
                inset 0 1px 0 rgba(255, 255, 255, 0.3);
}

button[kind="primary"]:active {
    transform: translateY(0) scale(0.98);
}

/* ========== CONTAINERS & CARDS ========== */
div[data-testid="stContainer"] {
    background: #ffffff;
    border-radius: 16px;
    padding: 25px;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.08);
    border: 1px solid rgba(0, 61, 130, 0.1);
}

/* Expander */
div[data-testid="stExpander"] {
    background: #ffffff;
    border: 2px solid #ffd4dd;
    border-radius: 12px;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

div[data-testid="stExpander"]:hover {
    border-color: #ff6b9d;
    box-shadow: 0 4px 15px rgba(255, 107, 157, 0.15);
}

/* ========== DATAFRAMES ========== */
div[data-testid="stDataFrame"] {
    border: 2px solid #e0e6ed;
    border-radius: 12px;
    overflow: hidden;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
}

/* ========== INFO/WARNING/ERROR BOXES ========== */
div[data-baseweb="notification"] {
    border-radius: 12px;
    border-left-width: 5px !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
    padding: 20px !important;
}

/* Info box */
div[data-baseweb="notification"][data-testid*="stInfo"] {
    background: linear-gradient(135deg, #ffe8f0 0%, #ffd4dd 100%);
    border-left-color: #ff6b9d !important;
}

/* Success box */
div[data-baseweb="notification"][data-testid*="stSuccess"] {
    background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
    border-left-color: #28a745 !important;
}

/* Warning box */
div[data-baseweb="notification"][data-testid*="stWarning"] {
    background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
    border-left-color: #ffc107 !important;
}

/* Error box */
div[data-baseweb="notification"][data-testid*="stError"] {
    background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%);
    border-left-color: #dc3545 !important;
}

/* ========== DIVIDER ========== */
hr {
    border: none;
    height: 3px;
    background: linear-gradient(90deg, transparent, #ffb3c6, transparent);
    margin: 30px 0;
}

/* ========== PROGRESS BAR ========== */
div[data-testid="stProgress"] > div {
    background: linear-gradient(90deg, #ff6b9d, #ff85a1, #ffb3c6);
    border-radius: 10px;
    box-shadow: 0 2px 10px rgba(255, 107, 157, 0.3);
}

/* ========== SPINNER ========== */
div[data-testid="stSpinner"] > div {
    border-top-color: #ffb3c6 !important;
}

/* ========== TOOLTIPS & CAPTIONS ========== */
.stCaption {
    color: #6b7280 !important;
    font-weight: 500 !important;
}

/* ========== RESPONSIVE ENHANCEMENTS ========== */
@media (max-width: 768px) {
    .banner-title-container {
        padding: 25px 20px;
    }

    .banner-title-container h1 {
        font-size: 2rem !important;
    }

    button[data-testid="stTab"] {
        padding: 10px 15px;
        font-size: 0.9rem;
    }
}

/* ========== SCROLL BAR ========== */
::-webkit-scrollbar {
    width: 10px;
    height: 10px;
}

::-webkit-scrollbar-track {
    background: #f5f7fa;
    border-radius: 10px;
}

::-webkit-scrollbar-thumb {
    background: linear-gradient(180deg, #ff6b9d, #ff85a1);
    border-radius: 10px;
}

::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(180deg, #e91e63, #f06292);
}


</style>
""", unsafe_allow_html=True)


# =========================
# HÀM GỌI GEMINI API (GIỮ NGUYÊN LOGIC)
# =========================

def get_ai_analysis(data_payload: dict, api_key: str) -> str:
    """
    Sử dụng Gemini API để phân tích chỉ số tài chính.
    """
    if not _GEMINI_OK:
        return "Lỗi: Thiếu thư viện google-genai (cần cài đặt: pip install google-genai)."

    client = genai.Client(api_key=api_key)

    sys_prompt = (
        "Bạn là chuyên gia phân tích tín dụng doanh nghiệp tại ngân hàng Việt Nam. "
        "Phân tích toàn diện dựa trên 14 chỉ số tài chính được cung cấp và PD (chủ yếu là PD cuối cùng của mô hình Stacking) . Lưu ý PD trong mô hình này được tính theo bối cảnh doanh nghiệp Việt Nam"
        "Nêu rõ: (1) Khả năng sinh lời, (2) Thanh khoản, (3) Cơ cấu nợ, (4) Hiệu quả hoạt động. "
        "Kết thúc bằng khuyến nghị in hoa: CHO VAY hoặc KHÔNG CHO VAY, kèm 2–3 điều kiện nếu CHO VAY. "
        "Viết bằng tiếng Việt súc tích, chuyên nghiệp."
    )

    # Gửi tên tiếng Việt dễ hiểu hơn cho AI
    user_prompt = "Bộ chỉ số tài chính và PD cần phân tích:\n" + str(data_payload) + "\n\nHãy phân tích và đưa ra khuyến nghị."

    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                {"role": "user", "parts": [{"text": sys_prompt + "\n\n" + user_prompt}]}
            ],
            config={"system_instruction": sys_prompt}
        )
        return response.text
    except APIError as e:
        return f"Lỗi gọi API Gemini: {e}"
    except Exception as e:
        return f"Lỗi không xác định: {e}"


def chat_with_gemini(user_message: str, api_key: str, context_data: dict = None) -> str:
    """
    Chatbot với Gemini AI để trả lời câu hỏi của người dùng về phân tích tín dụng.

    Args:
        user_message: Câu hỏi từ người dùng
        api_key: API key của Gemini
        context_data: Dữ liệu ngữ cảnh (chỉ số tài chính, PD, phân tích trước đó)

    Returns:
        Câu trả lời từ Gemini AI
    """
    if not _GEMINI_OK:
        return "Lỗi: Thiếu thư viện google-genai (cần cài đặt: pip install google-genai)."

    client = genai.Client(api_key=api_key)

    # System prompt cho chatbot
    sys_prompt = (
        "Bạn là chuyên gia tư vấn tín dụng doanh nghiệp tại ngân hàng. "
        "Nhiệm vụ của bạn là trả lời các câu hỏi của người dùng về phân tích tín dụng một cách chuyên nghiệp, "
        "dựa trên dữ liệu tài chính và phân tích đã được cung cấp. "
        "Trả lời súc tích, rõ ràng, dễ hiểu bằng tiếng Việt. "
        "Nếu cần, đưa ra các khuyến nghị hoặc giải thích chi tiết về các chỉ số tài chính."
    )

    # Tạo context prompt nếu có dữ liệu
    context_prompt = ""
    if context_data:
        context_prompt = "\n\nDữ liệu ngữ cảnh:\n" + str(context_data)

    full_prompt = user_message + context_prompt

    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                {"role": "user", "parts": [{"text": full_prompt}]}
            ],
            config={"system_instruction": sys_prompt}
        )
        return response.text
    except APIError as e:
        return f"Lỗi gọi API Gemini: {e}"
    except Exception as e:
        return f"Lỗi không xác định: {e}"


# =========================
# HÀM LẤY DỮ LIỆU TÀI CHÍNH TỰ ĐỘNG TỪ GEMINI API
# =========================

@st.cache_data(ttl=2592000)  # Cache 30 ngày (tự động cập nhật mỗi tháng)
def get_industry_data_from_ai(api_key: str, industry_name: str) -> dict:
    """
    Lấy dữ liệu ngành cụ thể từ Gemini API.

    Args:
        api_key: API key của Gemini
        industry_name: Tên ngành (VD: "Nông nghiệp", "Sản xuất", "Bất động sản"...)

    Returns:
        dict chứa dữ liệu ngành và phân tích
    """
    if not _GEMINI_OK:
        return None

    try:
        client = genai.Client(api_key=api_key)

        sys_prompt = """Bạn là chuyên gia phân tích kinh tế và dữ liệu ngành tại Việt Nam.
        Nhiệm vụ của bạn là cung cấp dữ liệu thống kê và phân tích về một ngành cụ thể."""

        user_prompt = f"""Hãy cung cấp dữ liệu và phân tích cho ngành **{industry_name}** tại Việt Nam trong 3 năm gần nhất.

        Trả về dữ liệu dưới dạng JSON với cấu trúc sau (CHỈ TRẢ VỀ JSON, KHÔNG GIẢI THÍCH):
        {{
            "industry_name": "{industry_name}",
            "revenue_growth_quarterly": {{
                "quarters": ["Q1-2022", "Q2-2022", ...],
                "growth_rate": [2.5, 3.1, ...]
            }},
            "avg_gross_margin_3y": 25.5,
            "avg_net_profit_margin": 8.3,
            "avg_debt_to_equity": 1.2,
            "pmi_monthly": {{
                "months": ["2024-01", "2024-02", ...],
                "pmi": [52.3, 51.8, ...]
            }},
            "new_vs_closed_businesses": {{
                "quarters": ["Q1-2022", "Q2-2022", ...],
                "new": [1200, 1350, ...],
                "closed": [450, 380, ...]
            }},
            "analysis": "Phân tích sơ bộ về tình hình ngành..."
        }}"""

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[{"role": "user", "parts": [{"text": sys_prompt + "\n\n" + user_prompt}]}],
            config={"system_instruction": sys_prompt}
        )

        import json
        import re

        response_text = response.text.strip()
        if "```json" in response_text:
            response_text = re.search(r'```json\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)
        elif "```" in response_text:
            response_text = re.search(r'```\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)

        data = json.loads(response_text)
        return data

    except Exception as e:
        st.error(f"Lỗi khi lấy dữ liệu ngành từ AI: {e}")
        return None


def get_macro_data_from_ai(api_key: str) -> dict:
    """
    Lấy dữ liệu vĩ mô nền kinh tế Việt Nam từ Gemini API.

    Returns:
        dict chứa dữ liệu vĩ mô và phân tích
    """
    if not _GEMINI_OK:
        return None

    try:
        client = genai.Client(api_key=api_key)

        sys_prompt = """Bạn là chuyên gia kinh tế vĩ mô Việt Nam.
        Nhiệm vụ của bạn là cung cấp dữ liệu vĩ mô quan trọng của nền kinh tế."""

        user_prompt = """Hãy cung cấp dữ liệu vĩ mô nền kinh tế Việt Nam trong 3-5 năm gần nhất.

        Trả về dữ liệu dưới dạng JSON với cấu trúc sau (CHỈ TRẢ VỀ JSON, KHÔNG GIẢI THÍCH):
        {
            "lending_rate_vs_interbank": {
                "quarters": ["Q1-2020", "Q2-2020", ...],
                "lending_rate": [8.5, 8.3, ...],
                "interbank_rate": [4.2, 4.0, ...]
            },
            "gdp_growth": {
                "quarters": ["Q1-2020", "Q2-2020", ...],
                "growth_rate": [3.7, 2.1, 6.7, 7.0, ...]
            },
            "unemployment_rate": {
                "years": ["2020", "2021", "2022", "2023", "2024"],
                "rate": [2.3, 2.5, 2.3, 2.2, 2.1]
            },
            "npl_ratio": {
                "quarters": ["Q1-2022", "Q2-2022", ...],
                "npl_rate": [1.9, 2.0, 2.1, ...],
                "default_rate": [0.5, 0.6, ...]
            },
            "financial_stress_index": {
                "months": ["2023-01", "2023-02", ...],
                "fsi": [0.3, 0.4, 0.2, ...]
            },
            "analysis": "Phân tích tổng quan về tình hình kinh tế vĩ mô..."
        }"""

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[{"role": "user", "parts": [{"text": sys_prompt + "\n\n" + user_prompt}]}],
            config={"system_instruction": sys_prompt}
        )

        import json
        import re

        response_text = response.text.strip()
        if "```json" in response_text:
            response_text = re.search(r'```json\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)
        elif "```" in response_text:
            response_text = re.search(r'```\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)

        data = json.loads(response_text)
        return data

    except Exception as e:
        st.error(f"Lỗi khi lấy dữ liệu vĩ mô từ AI: {e}")
        return None


def get_financial_data_from_ai(api_key: str) -> pd.DataFrame:
    """
    Tự động lấy dữ liệu tài chính doanh nghiệp Việt Nam từ Gemini API.
    Dữ liệu bao gồm: Doanh thu, Tổng tài sản, Lợi nhuận, Nợ phải trả, VCSH theo quý.

    Returns:
        pd.DataFrame: DataFrame chứa dữ liệu tài chính theo quý
    """
    if not _GEMINI_OK:
        return None

    try:
        client = genai.Client(api_key=api_key)

        # Lấy quý hiện tại
        current_date = datetime.now()
        current_year = current_date.year
        current_month = current_date.month
        current_quarter = (current_month - 1) // 3 + 1

        # Prompt yêu cầu Gemini cung cấp dữ liệu tài chính
        sys_prompt = """Bạn là chuyên gia kinh tế và dữ liệu thống kê về doanh nghiệp Việt Nam.
        Hãy cung cấp dữ liệu tài chính tổng hợp của khu vực doanh nghiệp Việt Nam theo quý,
        dựa trên các nguồn thống kê đáng tin cậy như GSO (Tổng cục Thống kê Việt Nam),
        Bộ Kế hoạch và Đầu tư, hoặc các báo cáo kinh tế vĩ mô.

        Trả về dữ liệu dưới dạng JSON với cấu trúc sau:
        {
            "quarters": ["Q1-2021", "Q2-2021", ...],
            "revenue": [số liệu doanh thu tỷ VNĐ, ...],
            "assets": [số liệu tổng tài sản tỷ VNĐ, ...],
            "profit": [số liệu lợi nhuận tỷ VNĐ, ...],
            "debt": [số liệu nợ phải trả tỷ VNĐ, ...],
            "equity": [số liệu VCSH tỷ VNĐ, ...]
        }

        Chỉ trả về JSON, không giải thích thêm."""

        user_prompt = f"""Hãy cung cấp dữ liệu tài chính tổng hợp của khu vực doanh nghiệp Việt Nam
        từ quý Q1-2021 đến quý Q{current_quarter}-{current_year}.

        Bao gồm các chỉ số:
        - Doanh thu (Revenue) - tổng doanh thu khu vực doanh nghiệp, đơn vị tỷ VNĐ
        - Tổng tài sản (Total Assets) - tổng tài sản khu vực doanh nghiệp, đơn vị tỷ VNĐ
        - Lợi nhuận (Profit) - lợi nhuận sau thuế, đơn vị tỷ VNĐ
        - Nợ phải trả (Debt) - tổng nợ phải trả, đơn vị tỷ VNĐ
        - Vốn chủ sở hữu (Equity/VCSH) - tổng VCSH, đơn vị tỷ VNĐ

        Dữ liệu phải phản ánh xu hướng tăng trưởng thực tế của nền kinh tế Việt Nam.
        Chỉ trả về JSON thuần, không markdown, không giải thích."""

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                {"role": "user", "parts": [{"text": sys_prompt + "\n\n" + user_prompt}]}
            ],
            config={"system_instruction": sys_prompt}
        )

        # Parse JSON response
        import json
        import re

        response_text = response.text.strip()

        # Loại bỏ markdown code block nếu có
        if "```json" in response_text:
            response_text = re.search(r'```json\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)
        elif "```" in response_text:
            response_text = re.search(r'```\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)

        data = json.loads(response_text)

        # Tạo DataFrame
        df = pd.DataFrame({
            'Quý': data.get('quarters', []),
            'Doanh thu (tỷ VNĐ)': data.get('revenue', []),
            'Tổng tài sản (tỷ VNĐ)': data.get('assets', []),
            'Lợi nhuận (tỷ VNĐ)': data.get('profit', []),
            'Nợ phải trả (tỷ VNĐ)': data.get('debt', []),
            'VCSH (tỷ VNĐ)': data.get('equity', [])
        })

        return df

    except Exception as e:
        st.error(f"Lỗi khi lấy dữ liệu từ AI: {e}")
        return None


# =========================
# TÍNH X1..X14 TỪ 3 SHEET (CDKT/BCTN/LCTT) - SỬ DỤNG TÊN TIẾNG VIỆT (GIỮ NGUYÊN)
# =========================

# Bảng ánh xạ Tên chỉ số tiếng Việt
COMPUTED_COLS = [
    "Biên Lợi nhuận Gộp (X1)", "Biên Lợi nhuận Tr.Thuế (X2)", "ROA Tr.Thuế (X3)", 
    "ROE Tr.Thuế (X4)", "Tỷ lệ Nợ/TTS (X5)", "Tỷ lệ Nợ/VCSH (X6)", 
    "Thanh toán Hiện hành (X7)", "Thanh toán Nhanh (X8)", "Khả năng Trả lãi (X9)", 
    "Khả năng Trả nợ Gốc (X10)", "Tỷ lệ Tiền/VCSH (X11)", "Vòng quay HTK (X12)", 
    "Kỳ thu tiền BQ (X13)", "Hiệu suất Tài sản (X14)"
]

# Alias các dòng quan trọng trong từng sheet (GIỮ NGUYÊN)
ALIAS_IS = {
    "doanh_thu_thuan": ["Doanh thu thuần", "Doanh thu bán hàng", "Doanh thu thuần về bán hàng và cung cấp dịch vụ"],
    "gia_von": ["Giá vốn hàng bán"],
    "loi_nhuan_gop": ["Lợi nhuận gộp"],
    "chi_phi_lai_vay": ["Chi phí lãi vay", "Chi phí tài chính (trong đó: chi phí lãi vay)"],
    "loi_nhuan_truoc_thue": ["Tổng lợi nhuận kế toán trước thuế", "Lợi nhuận trước thuế", "Lợi nhuận trước thuế thu nhập DN"],
}
ALIAS_BS = {
    "tong_tai_san": ["Tổng tài sản"],
    "von_chu_so_huu": ["Vốn chủ sở hữu", "Vốn CSH"],
    "no_phai_tra": ["Nợ phải trả"],
    "tai_san_ngan_han": ["Tài sản ngắn hạn"],
    "no_ngan_han": ["Nợ ngắn hạn"],
    "hang_ton_kho": ["Hàng tồn kho"],
    "tien_tdt": ["Tiền và các khoản tương đương tiền", "Tiền và tương đương tiền"],
    "phai_thu_kh": ["Phải thu ngắn hạn của khách hàng", "Phải thu khách hàng"],
    "no_dai_han_den_han": ["Nợ dài hạn đến hạn trả", "Nợ dài hạn đến hạn"],
}
ALIAS_CF = {
    "khau_hao": ["Khấu hao TSCĐ", "Khấu hao", "Chi phí khấu hao"],
}

def _pick_year_cols(df: pd.DataFrame):
    """Chọn 2 cột năm gần nhất từ sheet (ưu tiên cột có nhãn là năm)."""
    numeric_years = []
    for c in df.columns[1:]:
        try:
            y = int(float(str(c).strip()))
            if 1990 <= y <= 2100:
                numeric_years.append((y, c))
        except Exception:
            continue
    if numeric_years:
        numeric_years.sort(key=lambda x: x[0])
        return numeric_years[-2][1], numeric_years[-1][1]
    # fallback: 2 cột cuối
    cols = df.columns[-2:]
    return cols[0], cols[1]

def _get_row_vals(df: pd.DataFrame, aliases: list[str]):
    """Tìm dòng theo alias. Trả về (prev, cur) theo 2 cột năm gần nhất."""
    label_col = df.columns[0]
    prev_col, cur_col = _pick_year_cols(df)
    mask = False
    for alias in aliases:
        mask = mask | df[label_col].astype(str).str.contains(alias, case=False, na=False)
    rows = df[mask]
    if rows.empty:
        return np.nan, np.nan
    row = rows.iloc[0]

    def to_num(x):
        try:
            # Xóa dấu phẩy, khoảng trắng
            return float(str(x).replace(",", "").replace(" ", ""))
        except Exception:
            return np.nan

    return to_num(row[prev_col]), to_num(row[cur_col])

def compute_ratios_from_three_sheets(xlsx_file) -> pd.DataFrame:
    """Đọc 3 sheet CDKT/BCTN/LCTT và tính X1..X14 theo yêu cầu."""
    bs = pd.read_excel(xlsx_file, sheet_name="CDKT", engine="openpyxl")
    is_ = pd.read_excel(xlsx_file, sheet_name="BCTN", engine="openpyxl")
    cf = pd.read_excel(xlsx_file, sheet_name="LCTT", engine="openpyxl")

    # ---- Tính toán các biến số tài chính (GIỮ NGUYÊN CÁCH TÍNH)
    DTT_prev, DTT_cur         = _get_row_vals(is_, ALIAS_IS["doanh_thu_thuan"])
    GVHB_prev, GVHB_cur = _get_row_vals(is_, ALIAS_IS["gia_von"])
    LNG_prev, LNG_cur         = _get_row_vals(is_, ALIAS_IS["loi_nhuan_gop"])
    LNTT_prev, LNTT_cur = _get_row_vals(is_, ALIAS_IS["loi_nhuan_truoc_thue"])
    LV_prev, LV_cur           = _get_row_vals(is_, ALIAS_IS["chi_phi_lai_vay"])
    TTS_prev, TTS_cur           = _get_row_vals(bs, ALIAS_BS["tong_tai_san"])
    VCSH_prev, VCSH_cur         = _get_row_vals(bs, ALIAS_BS["von_chu_so_huu"])
    NPT_prev, NPT_cur           = _get_row_vals(bs, ALIAS_BS["no_phai_tra"])
    TSNH_prev, TSNH_cur         = _get_row_vals(bs, ALIAS_BS["tai_san_ngan_han"])
    NNH_prev, NNH_cur           = _get_row_vals(bs, ALIAS_BS["no_ngan_han"])
    HTK_prev, HTK_cur           = _get_row_vals(bs, ALIAS_BS["hang_ton_kho"])
    Tien_prev, Tien_cur         = _get_row_vals(bs, ALIAS_BS["tien_tdt"])
    KPT_prev, KPT_cur           = _get_row_vals(bs, ALIAS_BS["phai_thu_kh"])
    NDH_prev, NDH_cur           = _get_row_vals(bs, ALIAS_BS["no_dai_han_den_han"])
    KH_prev, KH_cur = _get_row_vals(cf, ALIAS_CF["khau_hao"])

    if pd.notna(GVHB_cur): GVHB_cur = abs(GVHB_cur)
    if pd.notna(LV_cur):      LV_cur     = abs(LV_cur)
    if pd.notna(KH_cur):      KH_cur     = abs(KH_cur)

    def avg(a, b):
        if pd.isna(a) and pd.isna(b): return np.nan
        if pd.isna(a): return b
        if pd.isna(b): return a
        return (a + b) / 2.0
    TTS_avg    = avg(TTS_cur,    TTS_prev)
    VCSH_avg = avg(VCSH_cur, VCSH_prev)
    HTK_avg    = avg(HTK_cur,    HTK_prev)
    KPT_avg    = avg(KPT_cur,    KPT_prev)

    EBIT_cur = (LNTT_cur + LV_cur) if (pd.notna(LNTT_cur) and pd.notna(LV_cur)) else np.nan
    NDH_cur = 0.0 if pd.isna(NDH_cur) else NDH_cur

    def div(a, b):
        return np.nan if (b is None or pd.isna(b) or b == 0) else a / b

    # ==== TÍNH X1..X14 ==== (GIỮ NGUYÊN CÔNG THỨC)
    X1  = div(LNG_cur, DTT_cur)
    X2  = div(LNTT_cur, DTT_cur)
    X3  = div(LNTT_cur, TTS_avg)
    X4  = div(LNTT_cur, VCSH_avg)
    X5  = div(NPT_cur,  TTS_cur)
    X6  = div(NPT_cur,  VCSH_cur)
    X7  = div(TSNH_cur, NNH_cur)
    X8  = div((TSNH_cur - HTK_cur) if pd.notna(TSNH_cur) and pd.notna(HTK_cur) else np.nan, NNH_cur)
    X9  = div(EBIT_cur, LV_cur)
    X10 = div((EBIT_cur + (KH_cur if pd.notna(KH_cur) else 0.0)), (LV_cur + NDH_cur) if pd.notna(LV_cur) else np.nan)
    X11 = div(Tien_cur, VCSH_cur)
    X12 = div(GVHB_cur, HTK_avg)
    turnover = div(DTT_cur, KPT_avg)
    X13 = div(365.0, turnover) if pd.notna(turnover) and turnover != 0 else np.nan
    X14 = div(DTT_cur, TTS_avg)

    # Khởi tạo DataFrame với tên cột tiếng Việt mới
    ratios = pd.DataFrame([[X1, X2, X3, X4, X5, X6, X7, X8, X9, X10, X11, X12, X13, X14]],
                          columns=COMPUTED_COLS)
                          
    # Thêm cột X_1..X_14 ẩn để phục vụ việc dự báo mô hình
    ratios[[f"X_{i}" for i in range(1, 15)]] = ratios.values
    return ratios

# =========================
# HÀM ĐỌC RSS FEED
# =========================

@st.cache_data(ttl=7200)  # Cache 120 phút = 7200 giây
def fetch_rss_feed(url, source_name):
    """
    Đọc RSS feed từ URL và trả về 5 bài mới nhất.

    Parameters:
    - url: Đường dẫn RSS feed
    - source_name: Tên nguồn tin

    Returns:
    - List của dict chứa {title, link, published}
    """
    if not _FEEDPARSER_OK:
        return [{"title": "⚠️ Thiếu thư viện feedparser", "link": "#", "published": ""}]

    try:
        feed = feedparser.parse(url)
        articles = []

        # Lấy 5 bài mới nhất
        for entry in feed.entries[:5]:
            title = entry.get('title', 'Không có tiêu đề')
            link = entry.get('link', '#')

            # Xử lý thời gian
            published = entry.get('published', '')
            if not published:
                published = entry.get('updated', '')

            # Parse thời gian nếu có
            pub_time = ""
            if published:
                try:
                    from dateutil import parser as date_parser
                    dt = date_parser.parse(published)
                    pub_time = dt.strftime('%d/%m/%Y %H:%M')
                except:
                    pub_time = published

            articles.append({
                'title': title,
                'link': link,
                'published': pub_time
            })

        return articles if articles else [{"title": "Không có bài viết mới", "link": "#", "published": ""}]

    except Exception as e:
        return [{"title": f"⚠️ Lỗi khi đọc RSS: {str(e)[:50]}", "link": "#", "published": ""}]

# =========================
# UI & TRAIN MODEL
# =========================
np.random.seed(0)

# ========================================
# PREMIUM BANKING HEADER
# ========================================
st.markdown('<div class="banner-title-container">', unsafe_allow_html=True)

# Thêm logo nếu có (optional)
col_logo, col_title = st.columns([1, 5])
with col_logo:
    try:
        st.image("logo-agribank.jpg", width=120)
    except:
        st.markdown("🏦")

with col_title:
    st.markdown("""
        <h1 style='margin: 0; padding: 0;'>CHƯƠNG TRÌNH ĐÁNH GIÁ RỦI RO TÍN DỤNG</h1>
        <h3 style='margin: 5px 0 0 0;'>Dự báo Xác suất Vỡ nợ KHDN (PD) & Phân tích AI Chuyên sâu</h3>
    """, unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# Thông báo hướng dẫn về sidebar
st.markdown("""
<div style='
    background: rgba(255, 182, 193, 0.15);
    padding: 12px 20px;
    border-radius: 10px;
    margin: 15px 0;
    border-left: 4px solid #ff6b9d;
'>
    <p style='margin: 0; color: #c2185b; font-size: 14px;'>
        <strong>💡 Note:</strong> Bấm vào nút <strong>mũi tên (&gt;&gt;)</strong> ở góc trái trên để <strong>ẩn/hiện tab tải file huấn luyện</strong>.
        Tab này chứa chức năng tải file CSV để xây dựng mô hình dự báo.
    </p>
</div>
""", unsafe_allow_html=True)

# Load dữ liệu huấn luyện (CSV có default, X_1..X_14) - Giữ nguyên logic load data
try:
    df = pd.read_csv('DATASET.csv', encoding='latin-1')
    # Tên cột cho việc huấn luyện (phải giữ nguyên X_1..X_14)
    MODEL_COLS = [f"X_{i}" for i in range(1, 15)]
except Exception:
    df = None

# ========================================
# SIDEBAR - HƯỚNG DẪN VÀ UPLOAD FILE
# ========================================

# Thêm header rõ ràng cho sidebar
st.sidebar.markdown("""
<div style='
    background: rgba(255, 255, 255, 0.1);
    padding: 15px;
    border-radius: 10px;
    margin-bottom: 20px;
    border-left: 4px solid #ffb3c6;
'>
    <h3 style='color: #ffffff; margin: 0; font-size: 18px;'>
        📁 TẢI DỮ LIỆU HUẤN LUYỆN
    </h3>
    <p style='color: #e8f4f8; margin: 8px 0 0 0; font-size: 13px;'>
        Tải file CSV để xây dựng mô hình dự báo
    </p>
</div>
""", unsafe_allow_html=True)

# Upload file
uploaded_file = st.sidebar.file_uploader("📂 Tải CSV Dữ liệu Huấn luyện", type=['csv'], label_visibility="collapsed")
if uploaded_file is not None:
    df = pd.read_csv(uploaded_file, encoding='latin-1')
    MODEL_COLS = [f"X_{i}" for i in range(1, 15)]
    
# Định nghĩa các Tabs
# ------------------------------------------------------------------------------------------------
# THAY ĐỔI 4: Vị trí Tabs được giữ nguyên, CSS mới sẽ đảm bảo Tabs có màu
# Tab mới: Dashboard tài chính doanh nghiệp (GSO) và Tin tức tài chính
# ------------------------------------------------------------------------------------------------
tab_predict, tab_dashboard, tab_news, tab_authors, tab_build, tab_goal = st.tabs([
    "🚀 Sử dụng mô hình dự báo",
    "📊 Dashboard tài chính doanh nghiệp",
    "📰 Tin tức tài chính",
    "👥 Nhóm tác giả",
    "🛠️ Xây dựng mô hình",
    "🎯 Mục tiêu của mô hình"
])

# --- Logic xử lý khi chưa có data huấn luyện ---
if df is None:
    st.sidebar.info("💡 Hãy tải file CSV huấn luyện (có cột 'default' và X_1...X_14) để xây dựng mô hình.")
    
    # Logic cho các tab khi thiếu data huấn luyện
    with tab_predict:
        st.header("⚡ Dự báo PD & Phân tích AI cho Hồ sơ mới")
        st.warning("⚠️ **Không thể dự báo PD**. Vui lòng tải file **CSV Dữ liệu Huấn luyện** ở sidebar để xây dựng mô hình Logistic Regression.")
        up_xlsx = st.file_uploader("Tải **ho_so_dn.xlsx**", type=["xlsx"], key="ho_so_dn")
        if up_xlsx is None:
            st.info("Hãy tải **ho_so_dn.xlsx** (đủ 3 sheet) để tính X1…X14 và phân tích AI.")

    with tab_goal:
        st.header("🎯 Mục tiêu của Mô hình")
        st.info("Ứng dụng này cần dữ liệu huấn luyện để bắt đầu hoạt động.")
    
    with tab_build:
          st.header("🛠️ Xây dựng & Đánh giá Mô hình LogReg")
          st.error("❌ **Không thể xây dựng mô hình**. Vui lòng tải file **CSV Dữ liệu Huấn luyện** ở sidebar để bắt đầu.")
          
    st.stop()
# ------------------------------------------------------------------------------------------------

# Hiển thị trạng thái thư viện AI (Sử dụng cột để bố trí đẹp hơn)
col_ai_status, col_date = st.columns([3, 1])
with col_ai_status:
    ai_status = ("✅ sẵn sàng (cần 'GEMINI_API_KEY' trong Secrets)" if _GEMINI_OK else "⚠️ Thiếu thư viện google-genai.")
    st.caption(f"🔎 Trạng thái Gemini AI: **<span style='color: #004c99; font-weight: bold;'>{ai_status}</span>**", unsafe_allow_html=True)
with col_date:
    st.caption(f"📅 Cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

st.divider()

# Kiểm tra cột cần thiết
required_cols = ['default'] + MODEL_COLS
missing = [c for c in required_cols if c not in df.columns]
if missing:
    st.error(f"❌ Thiếu cột: **{missing}**. Vui lòng kiểm tra lại file CSV huấn luyện.")
    st.stop()


# ================================================================================================
# NÂNG CẤP MÔ HÌNH: Từ Logistic đơn lẻ lên StackingClassifier với 3 base models
# ================================================================================================
X = df[MODEL_COLS] # Chỉ lấy các cột X_1..X_14
y = df['default'].astype(int)

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)

# Định nghĩa 3 Base Models
model_logistic = LogisticRegression(random_state=42, max_iter=1000, class_weight="balanced", solver="lbfgs")
model_rf = RandomForestClassifier(n_estimators=100, random_state=42, max_depth=10, class_weight="balanced")
model_xgb = XGBClassifier(n_estimators=100, random_state=42, max_depth=6, learning_rate=0.1,
                          use_label_encoder=False, eval_metric='logloss')

# Tạo StackingClassifier với LogisticRegression làm meta-model
estimators = [
    ('logistic', model_logistic),
    ('random_forest', model_rf),
    ('xgboost', model_xgb)
]
model = StackingClassifier(
    estimators=estimators,
    final_estimator=LogisticRegression(random_state=42, max_iter=1000),
    cv=5,  # Cross-validation 5-fold
    stack_method='predict_proba',  # Dùng probability để stack
    n_jobs=-1  # Sử dụng tất cả CPU cores
)

# Train tất cả models
model.fit(X_train, y_train)

# Dự báo & đánh giá cho Stacking Model (Model chính)
y_pred_in = model.predict(X_train)
y_proba_in = model.predict_proba(X_train)[:, 1]
y_pred_out = model.predict(X_test)
y_proba_out = model.predict_proba(X_test)[:, 1]

# Train riêng 3 base models để lấy PD riêng biệt (để hiển thị)
model_logistic.fit(X_train, y_train)
model_rf.fit(X_train, y_train)
model_xgb.fit(X_train, y_train)

# Tính PD từ 3 base models trên test set
y_proba_logistic_out = model_logistic.predict_proba(X_test)[:, 1]
y_proba_rf_out = model_rf.predict_proba(X_test)[:, 1]
y_proba_xgb_out = model_xgb.predict_proba(X_test)[:, 1]

metrics_in = {
    "accuracy_in": accuracy_score(y_train, y_pred_in),
    "precision_in": precision_score(y_train, y_pred_in, zero_division=0),
    "recall_in": recall_score(y_train, y_pred_in, zero_division=0),
    "f1_in": f1_score(y_train, y_pred_in, zero_division=0),
    "auc_in": roc_auc_score(y_train, y_proba_in),
}
metrics_out = {
    "accuracy_out": accuracy_score(y_test, y_pred_out),
    "precision_out": precision_score(y_test, y_pred_out, zero_division=0),
    "recall_out": recall_score(y_test, y_pred_out, zero_division=0),
    "f1_out": f1_score(y_test, y_pred_out, zero_division=0),
    "auc_out": roc_auc_score(y_test, y_proba_out),
}

# --- CÁC PHẦN UI DỰA TRÊN TABS ---

with tab_goal:
    st.header("🎯 Mục tiêu của Mô hình")
    st.markdown("""
    **Dự báo xác suất vỡ nợ (PD) của khách hàng doanh nghiệp** dựa trên bộ chỉ số $\\text{X1}–\\text{X14}$
    (tính từ Bảng Cân đối Kế toán, Báo cáo Kết quả Kinh doanh và Báo cáo Lưu chuyển Tiền tệ).

    **Mô hình Nâng cấp**: Sử dụng **Stacking Classifier** với 3 base models (Logistic + RandomForest + XGBoost)
    để đạt độ chính xác cao hơn và khả năng giải thích tốt hơn so với mô hình đơn lẻ.
    """)

    st.divider()

    # Mô tả về các biến đầu vào X1-X14
    st.markdown("### 📊 Các Biến Đầu vào (X1 - X14)")
    st.markdown("""
    Mô hình sử dụng **14 chỉ số tài chính** được tính toán từ 3 báo cáo tài chính chính của doanh nghiệp.
    Các chỉ số này phản ánh khả năng sinh lời, thanh khoản, cơ cấu nợ và hiệu quả hoạt động của doanh nghiệp.
    """)

    # Tạo 4 nhóm chỉ số
    st.markdown("#### 1. 💰 Nhóm Khả năng Sinh lời (Profitability)")
    st.markdown("""
    - **X1 - Biên Lợi nhuận Gộp**: Đo lường hiệu quả hoạt động kinh doanh cốt lõi
    - **X2 - Biên Lợi nhuận Trước Thuế**: Khả năng tạo lợi nhuận từ doanh thu
    - **X3 - ROA Trước Thuế**: Hiệu quả sử dụng tài sản để tạo lợi nhuận
    - **X4 - ROE Trước Thuế**: Khả năng sinh lời trên vốn chủ sở hữu
    """)

    st.markdown("#### 2. 🔒 Nhóm Cơ cấu Nợ & Đòn bẩy (Leverage)")
    st.markdown("""
    - **X5 - Tỷ lệ Nợ/Tổng Tài sản**: Mức độ sử dụng nợ trong cơ cấu tài sản
    - **X6 - Tỷ lệ Nợ/Vốn Chủ sở hữu**: Đo lường đòn bẩy tài chính và rủi ro vỡ nợ
    """)

    st.markdown("#### 3. 💧 Nhóm Thanh khoản (Liquidity)")
    st.markdown("""
    - **X7 - Thanh toán Hiện hành**: Khả năng trả nợ ngắn hạn bằng tài sản ngắn hạn
    - **X8 - Thanh toán Nhanh**: Khả năng thanh toán nợ ngắn hạn bằng tài sản có tính thanh khoản cao
    - **X9 - Khả năng Trả lãi**: Đo lường khả năng trả lãi vay từ lợi nhuận
    - **X10 - Khả năng Trả nợ Gốc**: Khả năng hoàn trả gốc vay từ dòng tiền hoạt động
    - **X11 - Tỷ lệ Tiền/Vốn Chủ sở hữu**: Lượng tiền mặt so với vốn chủ sở hữu
    """)

    st.markdown("#### 4. ⚙️ Nhóm Hiệu quả Hoạt động (Efficiency)")
    st.markdown("""
    - **X12 - Vòng quay Hàng tồn kho**: Tốc độ luân chuyển hàng tồn kho
    - **X13 - Kỳ thu tiền Bình quân**: Thời gian trung bình để thu hồi công nợ
    - **X14 - Hiệu suất Tài sản**: Khả năng tạo doanh thu từ tài sản
    """)

    st.divider()

    st.markdown("### 🎯 Cách Mô hình Dự báo")
    st.markdown("""
    1. **Thu thập dữ liệu**: Hệ thống đọc 3 báo cáo tài chính (BCĐKT, BCKQKD, BCLCTT) từ file Excel
    2. **Tính toán chỉ số**: Tự động tính toán 14 chỉ số tài chính (X1-X14) từ các báo cáo
    3. **Dự báo PD**:
       - 3 Base Models (Logistic, RandomForest, XGBoost) dự báo độc lập
       - Meta-Model (Logistic) tổng hợp kết quả từ 3 models để đưa ra dự báo cuối cùng
    4. **Phân loại Rating**: Dựa trên PD, hệ thống phân loại doanh nghiệp theo 5 cấp độ (AAA-AA, A-BBB, BB, B, CCC-D)
    5. **Phân tích AI**: Gemini AI phân tích sâu các chỉ số và đưa ra khuyến nghị tín dụng
    """)

    st.info("💡 **Lưu ý**: Tất cả 14 chỉ số đều được tính toán tự động. Bạn chỉ cần tải file Excel chứa 3 báo cáo tài chính.")


with tab_build:
    st.header("🛠️ Xây dựng & Đánh giá Mô hình Stacking Ensemble")
    st.info("**Mô hình Stacking Classifier** đã được huấn luyện với **3 Base Models** (Logistic, RandomForest, XGBoost) + **Meta-Model** (Logistic) trên **20% dữ liệu Test (chưa thấy)**.")

    # Thêm expander để giải thích về Stacking Model với diagram
    with st.expander("ℹ️ Giải thích về Mô hình Stacking", expanded=True):
        st.markdown("""
        **Stacking Classifier** là phương pháp ensemble learning cao cấp:

        - **3 Base Models (Mô hình cơ sở)**:
          - **Logistic Regression**: Mô hình tuyến tính, dễ giải thích
          - **Random Forest**: Mô hình cây quyết định, xử lý tốt các quan hệ phi tuyến
          - **XGBoost**: Mô hình gradient boosting, hiệu suất cao

        - **Meta-Model (Mô hình tổng hợp)**:
          - **Logistic Regression** học cách kết hợp dự đoán từ 3 base models
          - Sử dụng probability predictions từ 3 base models làm đầu vào
          - **Cross-validation 5-fold** để tránh overfitting

        **Ưu điểm**: Kết hợp điểm mạnh của nhiều thuật toán, độ chính xác cao hơn, robust hơn.
        """)

        st.markdown("---")
        st.markdown("### 📊 Sơ đồ Hoạt động của Stacking Model")

        # Tạo diagram minh họa bằng text/markdown
        st.markdown("""
        ```
        ┌─────────────────────────────────────────────────────────────────┐
        │                    DỮ LIỆU ĐẦU VÀO (X1-X14)                    │
        │              14 Chỉ số Tài chính của Doanh nghiệp              │
        └────────────────────────────┬────────────────────────────────────┘
                                     │
                    ┌────────────────┼────────────────┐
                    ▼                ▼                ▼
        ┌──────────────────┐ ┌──────────────┐ ┌──────────────┐
        │  BASE MODEL 1    │ │ BASE MODEL 2 │ │ BASE MODEL 3 │
        │   LOGISTIC       │ │ RANDOM FOREST│ │   XGBOOST    │
        │   REGRESSION     │ │              │ │              │
        └────────┬─────────┘ └──────┬───────┘ └──────┬───────┘
                 │                  │                 │
                 │ PD₁ = 12.5%      │ PD₂ = 15.3%    │ PD₃ = 14.1%
                 │                  │                 │
                 └──────────────────┼─────────────────┘
                                    ▼
                    ┌───────────────────────────────┐
                    │       META-MODEL              │
                    │   LOGISTIC REGRESSION         │
                    │  (Tổng hợp 3 dự báo trên)     │
                    └───────────────┬───────────────┘
                                    │
                                    ▼
                    ┌───────────────────────────────┐
                    │  KẾT QUẢ CUỐI CÙNG: PD = 14%  │
                    │   Rating: BB (Trung bình)     │
                    │   Cần theo dõi                │
                    └───────────────────────────────┘
        ```
        """)

        st.markdown("""
        **Quy trình hoạt động:**
        1. **Bước 1**: Dữ liệu X1-X14 được đưa vào 3 Base Models độc lập
        2. **Bước 2**: Mỗi Base Model đưa ra dự báo PD riêng (PD₁, PD₂, PD₃)
        3. **Bước 3**: Meta-Model nhận 3 dự báo này làm đầu vào
        4. **Bước 4**: Meta-Model kết hợp thông minh để đưa ra PD cuối cùng
        5. **Bước 5**: Hệ thống phân loại Rating dựa trên PD cuối cùng
        """)

        st.success("💡 **Lợi ích**: Stacking giúp cân bằng giữa các models, giảm bias và variance, tăng độ chính xác!")

    
    # Hiển thị Metrics quan trọng bằng st.metric
    st.subheader("1. Tổng quan Kết quả Đánh giá (Test Set)")
    col_acc, col_auc, col_f1 = st.columns(3)
    
    col_acc.metric(label="Độ chính xác (Accuracy)", value=f"{metrics_out['accuracy_out']:.2%}")
    # Đảm bảo logic delta vẫn đúng
    col_auc.metric(label="Diện tích dưới đường cong (AUC)", value=f"{metrics_out['auc_out']:.3f}", delta=f"{metrics_in['auc_in'] - metrics_out['auc_out']:.3f}", delta_color="inverse")
    col_f1.metric(label="Điểm F1-Score", value=f"{metrics_out['f1_out']:.3f}")
    
    st.divider()

    # Thống kê chi tiết & Biểu đồ
    st.subheader("2. Dữ liệu và Trực quan hóa")
    
    with st.expander("📊 Thống kê Mô tả và Dữ liệu Mẫu"):
        st.markdown("##### Thống kê Mô tả các biến $X_1..X_{14}$")
        st.dataframe(df[MODEL_COLS].describe().style.format("{:.4f}"))
        st.markdown("##### 6 Dòng dữ liệu huấn luyện mẫu (Đầu/Cuối)")
        st.dataframe(pd.concat([df.head(3), df.tail(3)]))

    st.markdown("##### Biểu đồ Phân tán (Scatter Plot) với Đường Hồi quy Logisitc")
    col = st.selectbox('🔍 Chọn biến X muốn vẽ', options=MODEL_COLS, index=0, key="select_build_col")
    
    # Biểu đồ Scatter Plot và Đường Hồi quy Logisitc (GIỮ NGUYÊN LOGIC, CẢI THIỆN MÀU SẮC)
    if col in df.columns:
        try:
            # Dùng Streamlit.pyplot với theme banking hiện đại
            fig, ax = plt.subplots(figsize=(12, 7))

            # Set background color
            fig.patch.set_facecolor('#f8f9fa')
            ax.set_facecolor('#ffffff')

            # Scatter plot với màu sắc pink rose theme
            sns.scatterplot(data=df, x=col, y='default', alpha=0.65, ax=ax, hue='default',
                          palette=['#ff6b9d', '#ffb3c6'], s=80, edgecolor='white', linewidth=0.5)

            # Vẽ đường logistic regression theo 1 biến
            x_range = np.linspace(df[col].min(), df[col].max(), 100).reshape(-1, 1)
            X_temp = df[[col]].copy()
            y_temp = df['default']
            lr_temp = LogisticRegression(max_iter=1000)
            lr_temp.fit(X_temp, y_temp)
            x_test = pd.DataFrame({col: x_range[:, 0]})
            y_curve = lr_temp.predict_proba(x_test)[:, 1]
            ax.plot(x_range, y_curve, color='#c2185b', linewidth=4, label='Đường LogReg',
                   linestyle='-', alpha=0.9)

            # Styling cho tiêu đề và labels
            ax.set_title(f'Quan hệ giữa {col} và Xác suất Vỡ nợ', fontsize=16, fontweight='bold', color='#c2185b', pad=20)
            ax.set_ylabel('Xác suất Default (0: Non-Default, 1: Default)', fontsize=13, fontweight='600', color='#4a5568')
            ax.set_xlabel(col, fontsize=13, fontweight='600', color='#4a5568')

            # Grid styling
            ax.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#d0d0d0')
            ax.spines['bottom'].set_color('#d0d0d0')

            # Legend styling
            legend = ax.legend(title='Default Status', title_fontsize=11, fontsize=10,
                             frameon=True, fancybox=True, shadow=True)
            legend.get_frame().set_facecolor('#f8f9fa')
            legend.get_frame().set_alpha(0.9)

            st.pyplot(fig)
            plt.close(fig)
        except Exception as e:
            st.error(f"Lỗi khi vẽ biểu đồ: {e}")
    else:
        st.warning("Biến không tồn tại trong dữ liệu.")
    
    st.divider()

    st.subheader("3. Ma trận Nhầm lẫn và Bảng Metrics Chi tiết")
    col_cm, col_metrics_table = st.columns(2)
    
    with col_cm:
        st.markdown("##### Ma trận Nhầm lẫn (Test Set)")
        cm = confusion_matrix(y_test, y_pred_out)

        # Tạo custom colormap cho pink rose theme
        from matplotlib.colors import LinearSegmentedColormap
        colors_pink = ['#fff5f7', '#ffe8f0', '#ffd4dd', '#ff85a1', '#ff6b9d']
        n_bins = 100
        cmap_pink = LinearSegmentedColormap.from_list('pink_rose', colors_pink, N=n_bins)

        disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=['Non-Default (0)', 'Default (1)'])
        fig2, ax = plt.subplots(figsize=(7, 7))
        fig2.patch.set_facecolor('#f8f9fa')

        disp.plot(ax=ax, cmap=cmap_pink, colorbar=True)

        # Styling
        ax.set_title('Ma trận Nhầm lẫn', fontsize=14, fontweight='bold', color='#c2185b', pad=15)
        ax.set_xlabel('Predicted Label', fontsize=12, fontweight='600', color='#4a5568')
        ax.set_ylabel('True Label', fontsize=12, fontweight='600', color='#4a5568')

        st.pyplot(fig2)
        plt.close(fig2)
        
    with col_metrics_table:
        st.markdown("##### Bảng Metrics Chi tiết")
        dt = pd.DataFrame({
            "Metric": ["Accuracy", "Precision", "Recall", "F1-Score", "AUC"],
            "Train Set": [metrics_in['accuracy_in'], metrics_in['precision_in'], metrics_in['recall_in'], metrics_in['f1_in'], metrics_in['auc_in']],
            "Test Set": [metrics_out['accuracy_out'], metrics_out['precision_out'], metrics_out['recall_out'], metrics_out['f1_out'], metrics_out['auc_out']],
        }).set_index("Metric")
        # Thêm styling để làm nổi bật kết quả tốt nhất
        def highlight_max(s):
            is_max = s == s.max()
            return ['background-color: #e0f0ff' if v else '' for v in is_max]

        st.dataframe(dt.style.format("{:.4f}").apply(highlight_max, axis=1), use_container_width=True)

with tab_predict:
    # Trang này được hiển thị mặc định
    st.header("⚡ Dự báo PD & Phân tích AI cho Hồ sơ mới")
    
    # Sử dụng st.container và st.expander để tổ chức khu vực upload
    input_container = st.container(border=True)
    with input_container:
        st.markdown("##### 📥 Tải lên Hồ sơ Doanh nghiệp (Excel)")
        st.caption("File phải có đủ **3 sheet**: **CDKT** (Bảng Cân đối Kế toán) ; **BCTN** (Báo cáo Kết quả Kinh doanh) ; **LCTT** (Báo cáo Lưu chuyển Tiền tệ).")
        up_xlsx = st.file_uploader("Tải **ho_so_dn.xlsx**", type=["xlsx"], key="ho_so_dn_main", label_visibility="collapsed")
    
    if up_xlsx is not None:
        # Tính X1..X14 từ 3 sheet (GIỮ NGUYÊN)
        try:
            # Hiển thị thanh tiến trình giả lập (thêm hiệu ứng động)
            with st.spinner('Đang đọc và xử lý dữ liệu tài chính...'):
                ratios_df = compute_ratios_from_three_sheets(up_xlsx)
            
            # Tách riêng 14 cột tiếng Việt (hiển thị) và 14 cột tiếng Anh (dự báo)
            # ratios_display là DataFrame 1 cột: Index (Tên chỉ số) | Giá trị
            ratios_display = ratios_df[COMPUTED_COLS].T.rename(columns={0: 'Giá trị'})
            ratios_predict = ratios_df[MODEL_COLS]
            
        except Exception as e:
            st.error(f"❌ Lỗi tính chỉ số tài chính: Vui lòng kiểm tra lại cấu trúc 3 sheet trong file Excel. Chi tiết lỗi: {e}")
            st.stop()

        st.divider()
        st.markdown("### 1. 🔢 Các Chỉ số Tài chính Đã tính")
        
        # Tạo payload data cho AI (Sử dụng tên tiếng Việt)
        data_for_ai = ratios_display.to_dict()['Giá trị']
        
        # ================================================================================================
        # DỰ BÁO PD TỪ 4 MODELS: 3 Base Models + 1 Stacking Model
        # ================================================================================================
        probs = np.nan
        preds = np.nan
        probs_logistic = np.nan
        probs_rf = np.nan
        probs_xgb = np.nan

        # Kiểm tra mô hình có sẵn sàng dự báo không (đã train và cột khớp)
        if set(X.columns) == set(ratios_predict.columns):
            try:
                # Đảm bảo thứ tự cột cho predict đúng như thứ tự cột huấn luyện
                X_new = ratios_predict[X.columns]

                # 1. PD từ Stacking Model (Model chính - kết quả cuối cùng)
                probs_array = model.predict_proba(X_new)[:, 1]
                probs = float(probs_array[0])
                preds = int(probs >= 0.15)

                # 2. PD từ 3 Base Models (để hiển thị riêng)
                probs_logistic = float(model_logistic.predict_proba(X_new)[:, 1][0])
                probs_rf = float(model_rf.predict_proba(X_new)[:, 1][0])
                probs_xgb = float(model_xgb.predict_proba(X_new)[:, 1][0])

                # Thêm PD vào payload AI (chỉ dùng PD từ Stacking - kết quả cuối cùng)
                data_for_ai['Xác suất Vỡ nợ (PD) - Stacking'] = probs
                data_for_ai['Xác suất Vỡ nợ (PD) - Logistic'] = probs_logistic
                data_for_ai['Xác suất Vỡ nợ (PD) - RandomForest'] = probs_rf
                data_for_ai['Xác suất Vỡ nợ (PD) - XGBoost'] = probs_xgb
                data_for_ai['Dự đoán PD'] = "Default (Vỡ nợ)" if preds == 1 else "Non-Default (Không vỡ nợ)"
            except Exception as e:
                # Nếu có lỗi dự báo, chỉ cảnh báo, không dừng app
                st.warning(f"Không dự báo được PD: {e}")
        
        # ================================================================================================
        # HIỂN THỊ 4 PD: Phần này đã được di chuyển xuống dưới phần "Giải thích về Biểu đồ"
        # ================================================================================================

        # Hiển thị Chỉ số Tài chính
        st.markdown("#### 📊 Chi tiết Chỉ số Tài chính")
        pd_col_1, pd_col_2 = st.columns(2) # Chia làm 2 cột cho ratios

        ratios_list = ratios_display.index.tolist()
        mid_point = len(ratios_list) // 2
        # ratios_display đã có cấu trúc đúng: Index (Tên biến) | Giá trị (Con số)
        ratios_part1 = ratios_display.iloc[:mid_point]
        ratios_part2 = ratios_display.iloc[mid_point:]

        # Hàm styling với màu sắc nhẹ nhàng, ngọt ngào hơn
        def color_ratios(val):
            """Ánh xạ màu dựa trên tên chỉ số và giá trị với palette màu pastel"""
            value = val.values[0]

            # Màu pastel nhẹ nhàng
            PASTEL_GREEN = '#d4edda'      # Xanh lá nhạt
            PASTEL_BLUE = '#d1ecf1'       # Xanh dương nhạt
            PASTEL_YELLOW = '#fff3cd'     # Vàng nhạt
            PASTEL_ORANGE = '#ffe8d9'     # Cam nhạt
            PASTEL_RED = '#f8d7da'        # Đỏ nhạt
            PASTEL_PURPLE = '#e7d9f5'     # Tím nhạt

            # Chỉ số Thanh khoản (X7, X8) - Quan trọng cho khả năng thanh toán
            if "Thanh toán" in val.name:
                if value < 1.0:
                    return [f'background-color: {PASTEL_RED}; color: #721c24; font-weight: 600;' for _ in val]  # Nguy hiểm
                elif value >= 2.0:
                    return [f'background-color: {PASTEL_GREEN}; color: #155724; font-weight: 600;' for _ in val]  # Rất tốt
                elif value >= 1.5:
                    return [f'background-color: {PASTEL_BLUE}; color: #0c5460; font-weight: 500;' for _ in val]  # Tốt
                else:
                    return [f'background-color: {PASTEL_YELLOW}; color: #856404; font-weight: 500;' for _ in val]  # Cảnh báo

            # Chỉ số Nợ (X5, X6) - Cơ cấu tài chính
            if "Tỷ lệ Nợ/" in val.name:
                if value > 2.0:
                    return [f'background-color: {PASTEL_RED}; color: #721c24; font-weight: 600;' for _ in val]  # Rủi ro cao
                elif value > 1.0:
                    return [f'background-color: {PASTEL_ORANGE}; color: #975a16; font-weight: 500;' for _ in val]  # Cảnh báo
                elif value < 0.5:
                    return [f'background-color: {PASTEL_GREEN}; color: #155724; font-weight: 600;' for _ in val]  # Rất tốt
                else:
                    return [f'background-color: {PASTEL_BLUE}; color: #0c5460; font-weight: 500;' for _ in val]  # Tốt

            # Chỉ số Sinh lời (X1, X2, X3, X4) - Hiệu quả kinh doanh
            if "Lợi nhuận" in val.name or "ROA" in val.name or "ROE" in val.name:
                if value <= 0:
                    return [f'background-color: {PASTEL_RED}; color: #721c24; font-weight: 600;' for _ in val]  # Lỗ
                elif value > 0.15:
                    return [f'background-color: {PASTEL_GREEN}; color: #155724; font-weight: 600;' for _ in val]  # Xuất sắc
                elif value > 0.08:
                    return [f'background-color: {PASTEL_BLUE}; color: #0c5460; font-weight: 500;' for _ in val]  # Tốt
                elif value > 0.03:
                    return [f'background-color: {PASTEL_YELLOW}; color: #856404; font-weight: 500;' for _ in val]  # Trung bình
                else:
                    return [f'background-color: {PASTEL_ORANGE}; color: #975a16; font-weight: 500;' for _ in val]  # Yếu

            # Các chỉ số khác - màu tím pastel nhẹ nhàng
            return [f'background-color: {PASTEL_PURPLE}; color: #5a395f; font-weight: 500;' for _ in val]

        with pd_col_1:
             # Đảm bảo hiển thị Tên biến | Giá trị
             st.markdown("##### **💰 Chỉ số Tài chính (Phần 1)**")
             st.dataframe(
                 ratios_part1.style.apply(color_ratios, axis=1).format("{:.4f}").set_properties(**{
                     'font-size': '14px',
                     'border-radius': '5px',
                     'padding': '8px'
                 }),
                 use_container_width=True
             )

        with pd_col_2:
            # Đảm bảo hiển thị Tên biến | Giá trị
            st.markdown("##### **📈 Chỉ số Tài chính (Phần 2)**")
            st.dataframe(
                ratios_part2.style.apply(color_ratios, axis=1).format("{:.4f}").set_properties(**{
                    'font-size': '14px',
                    'border-radius': '5px',
                    'padding': '8px'
                }),
                use_container_width=True
            )
        # ================================================================================================

        st.divider()

        # ========================================
        # THÊM BIỂU ĐỒ VISUALIZATION CHO CÁC CHỈ SỐ TÀI CHÍNH
        # ========================================
        st.markdown("### 2. 📊 Trực quan hóa Các Chỉ số Tài chính")

        # Tạo 2 cột cho 2 loại biểu đồ
        chart_col1, chart_col2 = st.columns(2)

        with chart_col1:
            st.markdown("#### 📈 Biểu đồ Cột - Giá trị các Chỉ số")
            # Tạo bar chart
            fig_bar, ax_bar = plt.subplots(figsize=(8, 10))
            fig_bar.patch.set_facecolor('#fff5f7')
            ax_bar.set_facecolor('#ffffff')

            # Chuẩn bị data cho bar chart
            indicators = ratios_display.index.tolist()
            values = ratios_display['Giá trị'].values

            # Tạo màu gradient cho các bars
            bar_colors = plt.cm.RdPu(np.linspace(0.3, 0.9, len(indicators)))

            # Vẽ horizontal bar chart
            bars = ax_bar.barh(indicators, values, color=bar_colors, edgecolor='white', linewidth=1.5)

            # Thêm giá trị vào cuối mỗi bar
            for i, (bar, val) in enumerate(zip(bars, values)):
                width = bar.get_width()
                ax_bar.text(width, bar.get_y() + bar.get_height()/2,
                           f' {val:.3f}', ha='left', va='center',
                           fontsize=9, fontweight='600', color='#c2185b')

            # Styling
            ax_bar.set_xlabel('Giá trị', fontsize=12, fontweight='600', color='#4a5568')
            ax_bar.set_title('Các Chỉ số Tài chính', fontsize=14, fontweight='bold', color='#c2185b', pad=15)
            ax_bar.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d', axis='x')
            ax_bar.spines['top'].set_visible(False)
            ax_bar.spines['right'].set_visible(False)
            ax_bar.spines['left'].set_color('#d0d0d0')
            ax_bar.spines['bottom'].set_color('#d0d0d0')

            # Điều chỉnh layout để labels không bị cắt
            plt.tight_layout()
            st.pyplot(fig_bar)
            plt.close(fig_bar)

        with chart_col2:
            st.markdown("#### 🎯 Biểu đồ Radar - Phân tích Đa chiều")
            # Tạo radar chart (spider chart)
            fig_radar = plt.figure(figsize=(10, 10))
            fig_radar.patch.set_facecolor('#fff5f7')
            ax_radar = fig_radar.add_subplot(111, projection='polar')

            # Chuẩn bị data cho radar chart
            # Normalize các giá trị về khoảng 0-1 để dễ visualize
            from sklearn.preprocessing import MinMaxScaler
            scaler = MinMaxScaler()
            normalized_values = scaler.fit_transform(values.reshape(-1, 1)).flatten()

            # Tạo các góc cho mỗi chỉ số
            angles = np.linspace(0, 2 * np.pi, len(indicators), endpoint=False).tolist()
            normalized_values = normalized_values.tolist()

            # Đóng vòng tròn
            angles += angles[:1]
            normalized_values += normalized_values[:1]

            # Vẽ radar chart
            ax_radar.plot(angles, normalized_values, 'o-', linewidth=2.5, color='#ff6b9d', label='Chỉ số')
            ax_radar.fill(angles, normalized_values, alpha=0.25, color='#ffb3c6')

            # Thêm labels
            ax_radar.set_xticks(angles[:-1])
            # Rút ngắn tên chỉ số để dễ đọc
            short_labels = [label.split('(')[0].strip()[:20] for label in indicators]
            ax_radar.set_xticklabels(short_labels, size=8, color='#4a5568', fontweight='600')

            # Styling
            ax_radar.set_ylim(0, 1)
            ax_radar.set_title('Phân tích Đa chiều các Chỉ số\n(Normalized 0-1)',
                              fontsize=14, fontweight='bold', color='#c2185b', pad=20)
            ax_radar.grid(True, alpha=0.3, linestyle='--', linewidth=0.8, color='#ff6b9d')
            ax_radar.set_facecolor('#ffffff')

            plt.tight_layout()
            st.pyplot(fig_radar)
            plt.close(fig_radar)

        # Thêm expander với thông tin bổ sung
        with st.expander("ℹ️ Giải thích về Biểu đồ"):
            st.markdown("""
            **Biểu đồ Cột (Bar Chart):**
            - Hiển thị giá trị thực tế của từng chỉ số tài chính
            - Màu sắc gradient từ nhạt đến đậm để dễ phân biệt
            - Giá trị cụ thể được hiển thị bên cạnh mỗi cột

            **Biểu đồ Radar (Spider Chart):**
            - Hiển thị cân bằng tổng thể giữa các chỉ số
            - Giá trị được chuẩn hóa về thang 0-1 để dễ so sánh
            - Diện tích vùng phủ thể hiện độ mạnh của các chỉ số
            - Hình dạng đều = tốt, hình dạng lệch = cần cân bằng
            """)

        st.divider()

        # ================================================================================================
        # HIỂN THỊ 4 PD: 3 PD từ Base Models ở trên + 1 PD cuối cùng từ Stacking ở dưới (KẾT QUẢ CHÍNH)
        # ================================================================================================

        st.markdown("### 2. 🎯 Dự báo Xác suất Vỡ nợ (PD) từ 4 Mô hình")

        # Hiển thị 3 PD từ Base Models trên 1 hàng
        st.markdown("##### 📊 Dự báo từ 3 Mô hình Cơ sở")
        pd_col_logistic, pd_col_rf, pd_col_xgb = st.columns(3)

        with pd_col_logistic:
            pd_value_log = f"{probs_logistic:.2%}" if pd.notna(probs_logistic) else "N/A"
            st.metric(
                label="**PD - Logistic**",
                value=pd_value_log,
                delta="⬆️ Cao" if pd.notna(probs_logistic) and probs_logistic >= 0.15 else "⬇️ Thấp",
                delta_color=("inverse" if pd.notna(probs_logistic) and probs_logistic >= 0.15 else "normal")
            )

        with pd_col_rf:
            pd_value_rf = f"{probs_rf:.2%}" if pd.notna(probs_rf) else "N/A"
            st.metric(
                label="**PD - RandomForest**",
                value=pd_value_rf,
                delta="⬆️ Cao" if pd.notna(probs_rf) and probs_rf >= 0.15 else "⬇️ Thấp",
                delta_color=("inverse" if pd.notna(probs_rf) and probs_rf >= 0.15 else "normal")
            )

        with pd_col_xgb:
            pd_value_xgb = f"{probs_xgb:.2%}" if pd.notna(probs_xgb) else "N/A"
            st.metric(
                label="**PD - XGBoost**",
                value=pd_value_xgb,
                delta="⬆️ Cao" if pd.notna(probs_xgb) and probs_xgb >= 0.15 else "⬇️ Thấp",
                delta_color=("inverse" if pd.notna(probs_xgb) and probs_xgb >= 0.15 else "normal")
            )

        # Hiển thị PD Stacking nổi bật ở dưới
        st.markdown("##### 🏆 KẾT QUẢ DỰ BÁO CUỐI CÙNG (STACKING MODEL)")

        # Tạo layout để thu nhỏ chiều ngang (chỉ chiếm 1/2 màn hình ở giữa)
        col_left, col_center, col_right = st.columns([1, 2, 1])

        with col_center:
            # Sử dụng hàm classify_pd để lấy thông tin phân loại
            pd_classification = classify_pd(probs)

            # Sử dụng markdown với style đặc biệt - Màu nhạt hơn
            pd_value_stacking = f"{probs:.2%}" if pd.notna(probs) else "N/A"

            # Tạo màu nhẹ nhàng hơn dựa trên risk level
            light_colors = {
                '#28a745': 'rgba(40, 167, 69, 0.15)',    # Xanh lá rất nhạt
                '#5cb85c': 'rgba(92, 184, 92, 0.15)',    # Xanh lá nhạt
                '#ffc107': 'rgba(255, 193, 7, 0.15)',    # Vàng nhạt
                '#fd7e14': 'rgba(253, 126, 20, 0.15)',   # Cam nhạt
                '#dc3545': 'rgba(220, 53, 69, 0.15)',    # Đỏ nhạt
                '#6c757d': 'rgba(108, 117, 125, 0.15)'   # Xám nhạt
            }
            bg_color = light_colors.get(pd_classification['color'], 'rgba(255, 255, 255, 0.15)')
            border_color = pd_classification['color']
            text_color = pd_classification['color']

            st.markdown(f"""
            <div style='
                background: {bg_color};
                border: 2px solid {border_color};
                border-radius: 15px;
                padding: 25px;
                text-align: center;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
                margin: 15px 0;
            '>
                <div style='font-size: 16px; font-weight: 600; color: {text_color}; margin-bottom: 12px;'>
                    🏆 XÁC SUẤT VỠ NỢ (PD) - STACKING MODEL
                </div>
                <div style='font-size: 42px; font-weight: 800; color: {text_color}; margin: 15px 0;'>
                    {pd_value_stacking}
                </div>
                <div style='font-size: 20px; font-weight: 600; color: {text_color}; margin: 8px 0;'>
                    Rating: {pd_classification['rating']}
                </div>
                <div style='font-size: 17px; font-weight: 500; color: {text_color}; background: rgba(0,0,0,0.03); padding: 8px; border-radius: 8px; margin: 8px 0;'>
                    {pd_classification['classification']} ({pd_classification['range']})
                </div>
                <div style='font-size: 14px; color: {text_color}; margin-top: 8px; font-style: italic; opacity: 0.9;'>
                    📊 {pd_classification['meaning']}
                </div>
                <div style='font-size: 12px; color: {text_color}; margin-top: 12px; font-style: italic; opacity: 0.8;'>
                    💡 AI sử dụng kết quả này để phân tích và đề xuất quyết định tín dụng
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.divider()

        # Khu vực Phân tích AI
        st.markdown("### 3. 🧠 Phân tích AI & Khuyến nghị Tín dụng")

        # Khởi tạo session_state cho phân tích AI
        if 'show_ai_analysis' not in st.session_state:
            st.session_state['show_ai_analysis'] = False
        if 'ai_analysis' not in st.session_state:
            st.session_state['ai_analysis'] = ''
        if 'chat_messages' not in st.session_state:
            st.session_state['chat_messages'] = []
        if 'ai_context_data' not in st.session_state:
            st.session_state['ai_context_data'] = {}

        ai_container = st.container(border=True)
        with ai_container:
            st.markdown("Sử dụng AI để phân tích toàn diện các chỉ số và đưa ra khuyến nghị chuyên nghiệp.")

            # Tạo 2 cột cho nút phân tích và nút ẩn
            col_btn1, col_btn2 = st.columns([3, 1])

            with col_btn1:
                analyze_button = st.button("✨ Yêu cầu AI Phân tích & Đề xuất", use_container_width=True, type="primary", key="analyze_ai_btn")

            with col_btn2:
                if st.session_state['show_ai_analysis']:
                    hide_button = st.button("🔽 Ẩn phân tích", use_container_width=True, key="hide_ai_btn")
                    if hide_button:
                        st.session_state['show_ai_analysis'] = False
                        st.session_state['chat_messages'] = []
                        st.rerun()

            # Xử lý khi người dùng click nút phân tích
            if analyze_button:
                # Kiểm tra API Key: ưu tiên lấy từ secrets
                api_key = st.secrets.get("GEMINI_API_KEY")

                if api_key:
                    # Thêm thanh tiến trình đẹp mắt
                    progress_bar = st.progress(0, text="Đang gửi dữ liệu và chờ Gemini phân tích...")
                    for percent_complete in range(100):
                        import time
                        time.sleep(0.01) # Giả lập thời gian xử lý
                        progress_bar.progress(percent_complete + 1, text=f"Đang gửi dữ liệu và chờ Gemini phân tích... {percent_complete+1}%")

                    ai_result = get_ai_analysis(data_for_ai, api_key)
                    progress_bar.empty() # Xóa thanh tiến trình

                    # Lưu kết quả vào session_state
                    st.session_state['ai_analysis'] = ai_result
                    st.session_state['show_ai_analysis'] = True
                    st.session_state['ai_context_data'] = data_for_ai
                    st.session_state['chat_messages'] = []  # Reset chat khi phân tích mới
                    st.rerun()
                else:
                    st.error("❌ **Lỗi Khóa API**: Không tìm thấy Khóa API. Vui lòng cấu hình Khóa **'GEMINI_API_KEY'** trong Streamlit Secrets.")

        # Hiển thị kết quả phân tích AI và chatbot nếu đã có phân tích
        if st.session_state['show_ai_analysis'] and st.session_state['ai_analysis']:
            ai_result = st.session_state['ai_analysis']

            st.markdown("---")
            st.markdown("**Kết quả Phân tích Chi tiết từ Gemini AI:**")

            if "KHÔNG CHO VAY" in ai_result.upper():
                st.error("🚨 **KHUYẾN NGHỊ CUỐI CÙNG: KHÔNG CHO VAY**")
                st.snow()
            elif "CHO VAY" in ai_result.upper():
                st.success("✅ **KHUYẾN NGHỊ CUỐI CÙNG: CHO VAY**")
                st.balloons()
            else:
                st.info("💡 **KHUYẾN NGHỊ CUỐI CÙNG**")

            st.info(ai_result)

            # ===== NÚT ĐIỀU HƯỚNG ĐẾN DASHBOARD =====
            st.markdown("---")
            st.markdown("""
            <style>
            .dashboard-nav-button {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white !important;
                padding: 1rem 2rem;
                border-radius: 10px;
                font-size: 1.1rem;
                font-weight: 600;
                text-align: center;
                margin: 1rem 0;
                box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                border: none;
                cursor: pointer;
                transition: all 0.3s ease;
            }
            .dashboard-nav-button:hover {
                transform: translateY(-2px);
                box-shadow: 0 6px 20px rgba(102, 126, 234, 0.6);
            }
            </style>
            """, unsafe_allow_html=True)

            # Khởi tạo session state cho tab navigation
            if 'navigate_to_dashboard' not in st.session_state:
                st.session_state['navigate_to_dashboard'] = False

            col_nav = st.columns([1, 2, 1])
            with col_nav[1]:
                if st.button("📊 Xem thêm DashBoard Tài chính Hỗ trợ Quyết định Cho vay",
                           use_container_width=True,
                           type="primary",
                           key="nav_to_dashboard_btn"):
                    st.session_state['navigate_to_dashboard'] = True
                    st.success("✅ Hãy chuyển sang tab **'📊 Dashboard tài chính doanh nghiệp'** ở phía trên để xem phân tích chi tiết!")
                    st.balloons()

            if st.session_state['navigate_to_dashboard']:
                st.info("💡 **Hướng dẫn**: Vui lòng click vào tab **'📊 Dashboard tài chính doanh nghiệp'** ở phía trên để xem thêm các biểu đồ và phân tích kinh tế hỗ trợ quyết định cho vay.")

            # ===== CHATBOT GEMINI AI =====
            st.markdown("---")
            st.markdown("#### 💬 Chatbot - Hỏi thêm thông tin")

            # Container cho chatbot
            chatbot_container = st.container(border=True)
            with chatbot_container:
                st.markdown("Bạn có thể hỏi thêm về kết quả phân tích, các chỉ số tài chính, hoặc bất kỳ câu hỏi nào liên quan đến tín dụng.")

                # Hiển thị lịch sử chat
                if st.session_state['chat_messages']:
                    st.markdown("**Lịch sử trò chuyện:**")
                    for msg in st.session_state['chat_messages']:
                        if msg['role'] == 'user':
                            st.markdown(f"**👤 Bạn:** {msg['content']}")
                        else:
                            st.markdown(f"**🤖 Gemini AI:** {msg['content']}")
                    st.markdown("---")

                # Form nhập câu hỏi
                with st.form(key='chat_form', clear_on_submit=True):
                    user_question = st.text_input(
                        "Nhập câu hỏi của bạn:",
                        placeholder="VD: Giải thích thêm về chỉ số thanh khoản...",
                        key='user_question_input'
                    )

                    col1, col2 = st.columns([1, 5])
                    with col1:
                        submit_button = st.form_submit_button("📤 Gửi", use_container_width=True)
                    with col2:
                        clear_button = st.form_submit_button("🗑️ Xóa lịch sử chat", use_container_width=True)

                # Xử lý khi người dùng gửi câu hỏi
                if submit_button and user_question.strip():
                    # Lấy API key
                    api_key = st.secrets.get("GEMINI_API_KEY")

                    # Lưu câu hỏi của user
                    st.session_state['chat_messages'].append({
                        'role': 'user',
                        'content': user_question
                    })

                    # Chuẩn bị context data cho chatbot
                    context_data = {
                        'chỉ_số_tài_chính': st.session_state.get('ai_context_data', data_for_ai),
                        'phân_tích_trước_đó': st.session_state['ai_analysis']
                    }

                    # Gọi chatbot API
                    with st.spinner("🤔 Gemini đang suy nghĩ..."):
                        bot_response = chat_with_gemini(user_question, api_key, context_data)

                    # Lưu response của bot
                    st.session_state['chat_messages'].append({
                        'role': 'assistant',
                        'content': bot_response
                    })

                    # Rerun để hiển thị tin nhắn mới
                    st.rerun()

                # Xử lý khi người dùng xóa lịch sử
                if clear_button:
                    st.session_state['chat_messages'] = []
                    st.rerun()

        st.divider()

        # ===== NÚT XUẤT FILE WORD =====
        st.markdown("### 4. 📄 Xuất Báo cáo Word")

        export_container = st.container(border=True)
        with export_container:
            st.markdown("Xuất toàn bộ phân tích (chỉ số tài chính, biểu đồ, PD, khuyến nghị AI) ra file Word chuyên nghiệp.")

            col_export1, col_export2 = st.columns([3, 1])

            with col_export1:
                company_name_input = st.text_input("Tên Khách hàng (tùy chọn):", value="KHÁCH HÀNG DOANH NGHIỆP", key="company_name_word")

            with col_export2:
                st.write("")  # Spacer

            if st.button("📥 Xuất file Word", use_container_width=True, type="primary", key="export_word_btn"):
                if not _WORD_OK:
                    st.error("❌ Thiếu thư viện python-docx. Không thể xuất Word.")
                else:
                    try:
                        with st.spinner("Đang tạo báo cáo Word..."):
                            # Lấy AI analysis từ session_state nếu có
                            ai_analysis_text = st.session_state.get('ai_analysis', '')

                            # Tạo lại figures để export (không hiển thị)
                            # Bar chart
                            fig_bar_export, ax_bar_export = plt.subplots(figsize=(8, 10))
                            fig_bar_export.patch.set_facecolor('#fff5f7')
                            ax_bar_export.set_facecolor('#ffffff')

                            indicators_export = ratios_display.index.tolist()
                            values_export = ratios_display['Giá trị'].values
                            bar_colors_export = plt.cm.RdPu(np.linspace(0.3, 0.9, len(indicators_export)))

                            bars_export = ax_bar_export.barh(indicators_export, values_export, color=bar_colors_export, edgecolor='white', linewidth=1.5)

                            for i, (bar, val) in enumerate(zip(bars_export, values_export)):
                                width = bar.get_width()
                                ax_bar_export.text(width, bar.get_y() + bar.get_height()/2,
                                           f' {val:.3f}', ha='left', va='center',
                                           fontsize=9, fontweight='600', color='#c2185b')

                            ax_bar_export.set_xlabel('Giá trị', fontsize=12, fontweight='600', color='#4a5568')
                            ax_bar_export.set_title('Các Chỉ số Tài chính', fontsize=14, fontweight='bold', color='#c2185b', pad=15)
                            ax_bar_export.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d', axis='x')
                            ax_bar_export.spines['top'].set_visible(False)
                            ax_bar_export.spines['right'].set_visible(False)
                            ax_bar_export.spines['left'].set_color('#d0d0d0')
                            ax_bar_export.spines['bottom'].set_color('#d0d0d0')
                            plt.tight_layout()

                            # Radar chart
                            fig_radar_export = plt.figure(figsize=(10, 10))
                            fig_radar_export.patch.set_facecolor('#fff5f7')
                            ax_radar_export = fig_radar_export.add_subplot(111, projection='polar')

                            from sklearn.preprocessing import MinMaxScaler
                            scaler_export = MinMaxScaler()
                            normalized_values_export = scaler_export.fit_transform(values_export.reshape(-1, 1)).flatten()

                            angles_export = np.linspace(0, 2 * np.pi, len(indicators_export), endpoint=False).tolist()
                            normalized_values_list_export = normalized_values_export.tolist()

                            angles_export += angles_export[:1]
                            normalized_values_list_export += normalized_values_list_export[:1]

                            ax_radar_export.plot(angles_export, normalized_values_list_export, 'o-', linewidth=2.5, color='#ff6b9d', label='Chỉ số')
                            ax_radar_export.fill(angles_export, normalized_values_list_export, alpha=0.25, color='#ffb3c6')

                            ax_radar_export.set_xticks(angles_export[:-1])
                            short_labels_export = [label.split('(')[0].strip()[:20] for label in indicators_export]
                            ax_radar_export.set_xticklabels(short_labels_export, size=8, color='#4a5568', fontweight='600')

                            ax_radar_export.set_ylim(0, 1)
                            ax_radar_export.set_title('Phân tích Đa chiều các Chỉ số\n(Normalized 0-1)',
                                              fontsize=14, fontweight='bold', color='#c2185b', pad=20)
                            ax_radar_export.grid(True, alpha=0.3, linestyle='--', linewidth=0.8, color='#ff6b9d')
                            ax_radar_export.set_facecolor('#ffffff')
                            plt.tight_layout()

                            # Tạo PD label
                            if pd.notna(probs) and pd.notna(preds):
                                pd_label_text = "Default (Vỡ nợ)" if preds == 1 else "Non-Default (Không vỡ nợ)"
                            else:
                                pd_label_text = "N/A"

                            # Generate Word
                            word_buffer = generate_word_report(
                                ratios_display=ratios_display,
                                pd_value=probs if pd.notna(probs) else np.nan,
                                pd_label=pd_label_text,
                                ai_analysis=ai_analysis_text,
                                fig_bar=fig_bar_export,
                                fig_radar=fig_radar_export,
                                company_name=company_name_input
                            )

                            # Close figures
                            plt.close(fig_bar_export)
                            plt.close(fig_radar_export)

                        st.success("✅ Báo cáo Word đã được tạo thành công!")

                        # Download button
                        st.download_button(
                            label="💾 Tải xuống Báo cáo Word",
                            data=word_buffer,
                            file_name=f"BaoCao_TinDung_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    except Exception as e:
                        st.error(f"❌ Lỗi khi tạo Word: {str(e)}")
                        st.exception(e)

    else:
        st.info("Hãy tải **ho_so_dn.xlsx** (đủ 3 sheet) để tính X1…X14, dự báo PD và phân tích AI.")

# ========================================
# TAB: DASHBOARD TÀI CHÍNH DOANH NGHIỆP
# ========================================
with tab_dashboard:
    st.header("📊 Dashboard Tài chính & Kinh tế")
    st.markdown("""
    Dashboard phân tích các chỉ số ngành và vĩ mô để hỗ trợ quyết định cho vay,
    dữ liệu được lấy tự động từ **Gemini AI** và các nguồn tin cậy.
    """)

    st.divider()

    # Hướng dẫn sử dụng
    info_container = st.container(border=True)
    with info_container:
        st.markdown("### 📖 Hướng dẫn sử dụng")
        st.info("""
        **Cách sử dụng Dashboard:**
        1. 📁 **Chọn loại phân tích**: Chọn ngành cụ thể hoặc "Tổng quan" để xem dữ liệu vĩ mô
        2. 🤖 **AI lấy dữ liệu tự động**: Bấm nút để Gemini AI lấy dữ liệu mới nhất
        3. 📊 **Xem biểu đồ**: Dữ liệu được hiển thị trực quan qua các biểu đồ
        4. 💡 **Đọc phân tích**: AI phân tích sơ bộ từng chỉ số
        5. 🔍 **Phân tích sâu**: Bấm nút để AI đánh giá ảnh hưởng đến quyết định cho vay
        """)

    st.divider()

    # Chọn loại phân tích: Ngành hoặc Tổng quan
    st.markdown("### 1️⃣ Chọn loại phân tích")

    # Danh sách ngành
    industries = [
        "Tổng quan (Vĩ mô)",
        "Nông nghiệp, Lâm nghiệp và Thủy sản",
        "Khai khoáng",
        "Công nghiệp chế biến, chế tạo",
        "Sản xuất và phân phối điện, khí đốt, nước",
        "Xây dựng",
        "Bán buôn và bán lẻ",
        "Vận tải và kho bãi",
        "Dịch vụ lưu trú và ăn uống",
        "Thông tin và truyền thông",
        "Hoạt động tài chính, ngân hàng và bảo hiểm",
        "Kinh doanh bất động sản",
        "Hoạt động chuyên môn, khoa học và công nghệ",
        "Giáo dục và đào tạo",
        "Y tế và hoạt động trợ giúp xã hội"
    ]

    selected_analysis = st.selectbox(
        "🔍 Chọn ngành hoặc tổng quan:",
        industries,
        index=0,
        key="analysis_type"
    )

    st.divider()

    # Nút lấy dữ liệu
    st.markdown("### 2️⃣ Lấy dữ liệu từ AI")
    get_data_btn = st.button("🤖 Lấy dữ liệu & Phân tích", use_container_width=True, type="primary")

    # Khởi tạo session_state cho cache
    if 'macro_data_cache' not in st.session_state:
        st.session_state['macro_data_cache'] = None
    if 'industry_data_cache' not in st.session_state:
        st.session_state['industry_data_cache'] = None
    if 'industry_selected_cache' not in st.session_state:
        st.session_state['industry_selected_cache'] = None
    if 'macro_analysis_result' not in st.session_state:
        st.session_state['macro_analysis_result'] = None
    if 'industry_analysis_result' not in st.session_state:
        st.session_state['industry_analysis_result'] = None

    # Xử lý khi người dùng bấm nút - CHỈ LẤY DỮ LIỆU
    if get_data_btn:
        if not _GEMINI_OK:
            st.error("❌ Thiếu thư viện google-genai. Vui lòng cài đặt: pip install google-genai")
        else:
            api_key = st.secrets.get("GEMINI_API_KEY")
            if not api_key:
                st.error("❌ **Lỗi Khóa API**: Không tìm thấy GEMINI_API_KEY trong Streamlit Secrets.")
            else:
                # Xác định loại phân tích
                is_macro = selected_analysis == "Tổng quan (Vĩ mô)"

                if is_macro:
                    # PHÂN TÍCH VĨ MÔ - CHỈ LẤY DỮ LIỆU
                    with st.spinner('🤖 Đang lấy dữ liệu vĩ mô từ Gemini AI...'):
                        macro_data = get_macro_data_from_ai(api_key)
                        st.session_state['macro_data_cache'] = macro_data

                    if macro_data:
                        st.success("✅ Đã lấy thành công dữ liệu vĩ mô!")
                    else:
                        st.error("⚠️ Không thể lấy dữ liệu vĩ mô từ AI.")

                else:
                    # PHÂN TÍCH NGÀNH - CHỈ LẤY DỮ LIỆU
                    # Kiểm tra xem ngành đã thay đổi chưa
                    if st.session_state['industry_selected_cache'] != selected_analysis:
                        with st.spinner(f'🤖 Đang lấy dữ liệu ngành "{selected_analysis}" từ Gemini AI...'):
                            industry_data = get_industry_data_from_ai(api_key, selected_analysis)
                            st.session_state['industry_data_cache'] = industry_data
                            st.session_state['industry_selected_cache'] = selected_analysis

                        if industry_data:
                            st.success(f"✅ Đã lấy thành công dữ liệu ngành {selected_analysis}!")
                        else:
                            st.error(f"⚠️ Không thể lấy dữ liệu ngành {selected_analysis} từ AI.")
                    else:
                        st.info(f"✅ Dữ liệu ngành {selected_analysis} đã có trong bộ nhớ!")

    st.divider()

    # ===== HIỂN THỊ DỮ LIỆU VÀ BIỂU ĐỒ (Chạy mỗi lần rerun) =====
    is_macro_selected = selected_analysis == "Tổng quan (Vĩ mô)"

    if is_macro_selected:
        # HIỂN THỊ DỮ LIỆU VĨ MÔ
        macro_data = st.session_state.get('macro_data_cache')
        if macro_data:
            st.markdown("### 📊 DỮ LIỆU VĨ MÔ NỀN KINH TẾ VIỆT NAM")

            # Hiển thị phân tích tổng quan
            if 'analysis' in macro_data:
                with st.expander("📝 Phân tích Tổng quan", expanded=True):
                    st.markdown(macro_data['analysis'])

            # 1. Lãi suất cho vay vs liên ngân hàng
            if 'lending_rate_vs_interbank' in macro_data:
                st.markdown("#### 💰 Lãi suất Cho vay & Liên ngân hàng")
                data = macro_data['lending_rate_vs_interbank']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                ax.plot(data['quarters'], data['lending_rate'], marker='o', linewidth=2.5,
                       markersize=7, color='#ff6b9d', label='Lãi suất cho vay', alpha=0.9)
                ax.plot(data['quarters'], data['interbank_rate'], marker='s', linewidth=2.5,
                       markersize=7, color='#4a90e2', label='Lãi suất liên ngân hàng', alpha=0.9)

                ax.set_xlabel('Quý', fontsize=13, fontweight='600')
                ax.set_ylabel('Lãi suất (%)', fontsize=13, fontweight='600')
                ax.set_title('Lãi suất Cho vay & Liên ngân hàng theo Quý', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--')
                ax.legend(fontsize=11)
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: Chênh lệch lãi suất cho vay và liên ngân hàng phản ánh mức độ rủi ro
                và biên lợi nhuận của ngân hàng. Xu hướng tăng/giảm ảnh hưởng đến chi phí vốn của doanh nghiệp.
                """)
                st.divider()

            # 2. Tăng trưởng GDP
            if 'gdp_growth' in macro_data:
                st.markdown("#### 📈 Tăng trưởng GDP")
                data = macro_data['gdp_growth']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                ax.bar(data['quarters'], data['growth_rate'], color='#50c878', alpha=0.8, edgecolor='white', linewidth=1.5)
                ax.axhline(y=0, color='red', linestyle='--', linewidth=1)

                ax.set_xlabel('Quý', fontsize=13, fontweight='600')
                ax.set_ylabel('Tăng trưởng GDP (%)', fontsize=13, fontweight='600')
                ax.set_title('Tăng trưởng GDP theo Quý', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--', axis='y')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: GDP tăng trưởng mạnh cho thấy nền kinh tế phát triển tốt,
                doanh nghiệp có nhiều cơ hội kinh doanh, tăng khả năng trả nợ.
                """)
                st.divider()

            # 3. Tỷ lệ thất nghiệp
            if 'unemployment_rate' in macro_data:
                st.markdown("#### 👥 Tỷ lệ Thất nghiệp")
                data = macro_data['unemployment_rate']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                ax.plot(data['years'], data['rate'], marker='o', linewidth=3,
                       markersize=8, color='#ffa500', alpha=0.9)
                ax.fill_between(data['years'], data['rate'], alpha=0.2, color='#ffa500')

                ax.set_xlabel('Năm', fontsize=13, fontweight='600')
                ax.set_ylabel('Tỷ lệ thất nghiệp (%)', fontsize=13, fontweight='600')
                ax.set_title('Tỷ lệ Thất nghiệp theo Năm', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: Tỷ lệ thất nghiệp thấp cho thấy thị trường lao động tốt,
                thu nhập ổn định, giảm rủi ro tín dụng cho cả doanh nghiệp và cá nhân.
                """)
                st.divider()

            # 4. Tỷ lệ nợ xấu
            if 'npl_ratio' in macro_data:
                st.markdown("#### ⚠️ Tỷ lệ Nợ xấu & Vỡ nợ")
                data = macro_data['npl_ratio']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                ax.plot(data['quarters'], data['npl_rate'], marker='o', linewidth=2.5,
                       markersize=7, color='#dc3545', label='Tỷ lệ nợ xấu', alpha=0.9)
                ax.plot(data['quarters'], data['default_rate'], marker='s', linewidth=2.5,
                       markersize=7, color='#ff6b9d', label='Tỷ lệ vỡ nợ', alpha=0.9)

                ax.set_xlabel('Quý', fontsize=13, fontweight='600')
                ax.set_ylabel('Tỷ lệ (%)', fontsize=13, fontweight='600')
                ax.set_title('Tỷ lệ Nợ xấu & Vỡ nợ Hệ thống Ngân hàng VN', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--')
                ax.legend(fontsize=11)
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: Tỷ lệ nợ xấu và vỡ nợ cao cảnh báo rủi ro tín dụng gia tăng trong hệ thống,
                cần thắt chặt tiêu chuẩn cho vay và tăng cường thẩm định.
                """)
                st.divider()

            # 5. Chỉ số căng thẳng tài chính
            if 'financial_stress_index' in macro_data:
                st.markdown("#### 📉 Chỉ số Căng thẳng Tài chính (FSI)")
                data = macro_data['financial_stress_index']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                colors = ['#28a745' if x < 0.5 else '#ffc107' if x < 0.7 else '#dc3545' for x in data['fsi']]
                ax.bar(data['months'], data['fsi'], color=colors, alpha=0.8, edgecolor='white', linewidth=1.5)
                ax.axhline(y=0.5, color='orange', linestyle='--', linewidth=1, label='Ngưỡng cảnh báo')
                ax.axhline(y=0.7, color='red', linestyle='--', linewidth=1, label='Ngưỡng nguy hiểm')

                ax.set_xlabel('Tháng', fontsize=13, fontweight='600')
                ax.set_ylabel('FSI', fontsize=13, fontweight='600')
                ax.set_title('Chỉ số Căng thẳng Tài chính theo Tháng', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--', axis='y')
                ax.legend(fontsize=11)
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: FSI đo lường mức độ căng thẳng trong hệ thống tài chính.
                FSI cao (>0.7) cảnh báo khủng hoảng, cần thận trọng khi cho vay.
                """)
                st.divider()

            # Nút phân tích sâu
            st.markdown("### 🔍 Phân tích Sâu bằng AI")
            analyze_macro_btn = st.button("💡 Phân tích ảnh hưởng đến Quyết định Cho vay",
                                         use_container_width=True, type="primary", key="analyze_macro")

            if analyze_macro_btn:
                api_key = st.secrets.get("GEMINI_API_KEY")
                if api_key:
                    with st.spinner('AI đang phân tích...'):
                        client = genai.Client(api_key=api_key)
                        prompt = f"""Dựa trên dữ liệu vĩ mô sau của nền kinh tế Việt Nam:
{macro_data}

Hãy phân tích CHI TIẾT ảnh hưởng của các chỉ số này đến quyết định cho vay của ngân hàng:
1. Rủi ro tín dụng tăng hay giảm?
2. Nên thắt chặt hay nới lỏng tiêu chuẩn cho vay?
3. Ngành nào nên ưu tiên cho vay, ngành nào nên hạn chế?
4. Khuyến nghị cụ thể cho chiến lược tín dụng.

Trả lời bằng tiếng Việt, có cấu trúc rõ ràng với các điểm bullet."""

                        response = client.models.generate_content(
                            model=MODEL_NAME,
                            contents=[{"role": "user", "parts": [{"text": prompt}]}]
                        )

                        st.session_state['macro_analysis_result'] = response.text
                else:
                    st.error("❌ Không tìm thấy GEMINI_API_KEY trong Streamlit Secrets.")

            # Hiển thị kết quả phân tích nếu có
            if st.session_state['macro_analysis_result']:
                st.markdown("---")
                st.markdown("#### 📊 Phân tích AI - Ảnh hưởng đến Quyết định Cho vay")
                st.success(st.session_state['macro_analysis_result'])
        else:
            st.info("💡 Hãy bấm nút '🤖 Lấy dữ liệu & Phân tích' để tải dữ liệu vĩ mô")

    else:
        # HIỂN THỊ DỮ LIỆU NGÀNH
        industry_data = st.session_state.get('industry_data_cache')
        if industry_data and st.session_state.get('industry_selected_cache') == selected_analysis:
            st.markdown(f"### 📊 DỮ LIỆU NGÀNH: {selected_analysis.upper()}")

            # Hiển thị phân tích sơ bộ
            if 'analysis' in industry_data:
                with st.expander("📝 Phân tích Sơ bộ", expanded=True):
                    st.markdown(industry_data['analysis'])

            # 1. Tốc độ tăng trưởng doanh thu
            if 'revenue_growth_quarterly' in industry_data:
                st.markdown("#### 💰 Tốc độ Tăng trưởng Doanh thu")
                data = industry_data['revenue_growth_quarterly']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                ax.plot(data['quarters'], data['growth_rate'], marker='o', linewidth=3,
                       markersize=8, color='#ff6b9d', alpha=0.9)
                ax.fill_between(data['quarters'], data['growth_rate'], alpha=0.2, color='#ffb3c6')
                ax.axhline(y=0, color='red', linestyle='--', linewidth=1)

                ax.set_xlabel('Quý', fontsize=13, fontweight='600')
                ax.set_ylabel('Tăng trưởng (%)', fontsize=13, fontweight='600')
                ax.set_title(f'Tốc độ Tăng trưởng Doanh thu - {selected_analysis}', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: Tăng trưởng doanh thu dương cho thấy ngành đang phát triển,
                doanh nghiệp trong ngành có khả năng trả nợ tốt hơn.
                """)
                st.divider()

            # 2. Biên lợi nhuận gộp và ròng
            st.markdown("#### 📊 Biên Lợi nhuận Trung bình Ngành")
            col1, col2, col3 = st.columns(3)

            with col1:
                if 'avg_gross_margin_3y' in industry_data:
                    st.metric("Biên LN Gộp TB (3 năm)", f"{industry_data['avg_gross_margin_3y']:.1f}%")

            with col2:
                if 'avg_net_profit_margin' in industry_data:
                    st.metric("Biên LN Ròng TB", f"{industry_data['avg_net_profit_margin']:.1f}%")

            with col3:
                if 'avg_debt_to_equity' in industry_data:
                    st.metric("Tỷ lệ Nợ/VCSH TB", f"{industry_data['avg_debt_to_equity']:.2f}")

            st.markdown("""
            **💡 Phân tích**: Biên lợi nhuận cao cho thấy ngành có khả năng sinh lời tốt.
            Tỷ lệ nợ/VCSH thấp (<1.5) là dấu hiệu tốt về cấu trúc vốn.
            """)
            st.divider()

            # 3. PMI ngành
            if 'pmi_monthly' in industry_data:
                st.markdown("#### 📈 Chỉ số PMI Ngành")
                data = industry_data['pmi_monthly']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                colors = ['#28a745' if x >= 50 else '#dc3545' for x in data['pmi']]
                ax.bar(data['months'], data['pmi'], color=colors, alpha=0.8, edgecolor='white', linewidth=1.5)
                ax.axhline(y=50, color='black', linestyle='--', linewidth=2, label='Ngưỡng 50')

                ax.set_xlabel('Tháng', fontsize=13, fontweight='600')
                ax.set_ylabel('PMI', fontsize=13, fontweight='600')
                ax.set_title(f'Chỉ số PMI - {selected_analysis}', fontsize=16, fontweight='bold', color='#c2185b')
                ax.grid(True, alpha=0.2, linestyle='--', axis='y')
                ax.legend(fontsize=11)
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: PMI >50 cho thấy ngành đang mở rộng, <50 cho thấy co hẹp.
                Xu hướng PMI giúp dự đoán sức khỏe ngành trong tương lai.
                """)
                st.divider()

            # 4. Doanh nghiệp mới vs giải thể
            if 'new_vs_closed_businesses' in industry_data:
                st.markdown("#### 🏢 Doanh nghiệp Đăng ký Mới vs Giải thể")
                data = industry_data['new_vs_closed_businesses']

                fig, ax = plt.subplots(figsize=(14, 6))
                fig.patch.set_facecolor('#fff5f7')
                ax.set_facecolor('#ffffff')

                x = np.arange(len(data['quarters']))
                width = 0.35

                ax.bar(x - width/2, data['new'], width, label='Đăng ký mới', color='#28a745', alpha=0.8)
                ax.bar(x + width/2, data['closed'], width, label='Giải thể', color='#dc3545', alpha=0.8)

                ax.set_xlabel('Quý', fontsize=13, fontweight='600')
                ax.set_ylabel('Số lượng DN', fontsize=13, fontweight='600')
                ax.set_title(f'DN Đăng ký Mới vs Giải thể - {selected_analysis}', fontsize=16, fontweight='bold', color='#c2185b')
                ax.set_xticks(x)
                ax.set_xticklabels(data['quarters'], rotation=45, ha='right')
                ax.legend(fontsize=11)
                ax.grid(True, alpha=0.2, linestyle='--', axis='y')
                plt.tight_layout()
                st.pyplot(fig)
                plt.close(fig)

                st.markdown("""
                **💡 Phân tích**: Số DN đăng ký mới > Giải thể cho thấy ngành đang hấp dẫn.
                Tỷ lệ giải thể cao cảnh báo rủi ro ngành đang gặp khó khăn.
                """)
                st.divider()

            # Nút phân tích sâu
            st.markdown("### 🔍 Phân tích Sâu bằng AI")
            analyze_industry_btn = st.button("💡 Phân tích ảnh hưởng đến Quyết định Cho vay",
                                use_container_width=True, type="primary", key="analyze_industry")

            if analyze_industry_btn:
                api_key = st.secrets.get("GEMINI_API_KEY")
                if api_key:
                    with st.spinner('AI đang phân tích...'):
                        client = genai.Client(api_key=api_key)
                        prompt = f"""Dựa trên dữ liệu ngành {selected_analysis} sau:
{industry_data}

Hãy phân tích CHI TIẾT:
1. Đánh giá tổng quan sức khỏe ngành này
2. Rủi ro tín dụng khi cho vay doanh nghiệp trong ngành
3. Các chỉ số đáng lo ngại và đáng mừng
4. Khuyến nghị CHO VAY hay KHÔNG CHO VAY cho ngành này, và các điều kiện cụ thể
5. Mức lãi suất và thời hạn cho vay phù hợp

Trả lời bằng tiếng Việt, có cấu trúc rõ ràng với các điểm bullet."""

                        response = client.models.generate_content(
                            model=MODEL_NAME,
                            contents=[{"role": "user", "parts": [{"text": prompt}]}]
                        )

                        st.session_state['industry_analysis_result'] = response.text
                else:
                    st.error("❌ Không tìm thấy GEMINI_API_KEY trong Streamlit Secrets.")

            # Hiển thị kết quả phân tích nếu có
            if st.session_state['industry_analysis_result']:
                st.markdown("---")
                st.markdown("#### 📊 Phân tích AI - Quyết định Cho vay")
                st.success(st.session_state['industry_analysis_result'])
        else:
            st.info(f"💡 Hãy bấm nút '🤖 Lấy dữ liệu & Phân tích' để tải dữ liệu ngành {selected_analysis}")


# ========================================
# TAB: TIN TỨC TÀI CHÍNH
# ========================================
with tab_news:
    st.header("📰 Tin tức Tài chính")
    st.markdown("""
    Tin tức tài chính mới nhất từ các nguồn uy tín tại Việt Nam.
    Dữ liệu tự động cập nhật mỗi **120 phút**.
    """)

    st.divider()

    if not _FEEDPARSER_OK:
        st.error("⚠️ **Thiếu thư viện feedparser**. Vui lòng cài đặt: `pip install feedparser python-dateutil`")
    else:
        # Định nghĩa các nguồn RSS
        rss_sources = {
            "📊 CafeF": "https://cafef.vn/thi-truong-chung-khoan.rss",
            "💼 Vietstock": "https://vietstock.vn/rss/tai-chinh.rss",
            "💰 Báo Đầu tư": "https://baodautu.vn/rss/kinh-doanh.rss",
            "🏢 VNExpress Kinh doanh": "https://vnexpress.net/rss/kinh-doanh.rss"
        }

        # Hiển thị thời gian cập nhật
        col_update, col_cache = st.columns([3, 1])
        with col_update:
            st.caption(f"🕐 Cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        with col_cache:
            st.caption("♻️ Cache: 120 phút")

        st.divider()

        # Tạo layout 2 cột
        col1, col2 = st.columns(2)

        sources_list = list(rss_sources.items())

        # Hiển thị nguồn tin 1 và 2 ở cột trái
        with col1:
            # Nguồn 1: CafeF
            source_name, source_url = sources_list[0]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # Nguồn 3: Báo Đầu tư
            source_name, source_url = sources_list[2]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

        # Hiển thị nguồn tin 2 và 4 ở cột phải
        with col2:
            # Nguồn 2: Vietstock
            source_name, source_url = sources_list[1]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # Nguồn 4: VNExpress
            source_name, source_url = sources_list[3]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

# ========================================
# TAB: NHÓM TÁC GIẢ
# ========================================
with tab_authors:
    # Header với hiệu ứng gradient
    st.markdown("""
        <div style='text-align: center; padding: 30px; background: linear-gradient(135deg, #fbc2eb 0%, #a6c1ee 100%); border-radius: 15px; margin-bottom: 30px; box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);'>
            <h1 style='color: white; margin: 0; font-size: 2.5rem; font-weight: 700; text-shadow: 2px 2px 4px rgba(0,0,0,0.2);'>
                👥 NHÓM ÁNH SÁNG SỐ
            </h1>
            <p style='color: #f0f0f0; font-size: 1.1rem; margin-top: 10px; font-weight: 500;'>
                Cuộc thi Agribank làm chủ công nghệ trong kỷ nguyên số 2025
            </p>
        </div>
    """, unsafe_allow_html=True)

    # Ảnh nhóm ở giữa
    col_left, col_center, col_right = st.columns([1, 2, 1])
    with col_center:
        try:
            st.image("NHOM ANH SANG SO.jpg", use_container_width=True, caption="Team Ánh Sáng Số - Ánh sáng của đổi mới, bước đi của tương lai")
        except:
            st.info("📸 Ảnh nhóm: NHOM ANH SANG SO.jpg")

    st.markdown("<br>", unsafe_allow_html=True)

    # Giới thiệu chung
    st.markdown("""
        <div style='text-align: center; padding: 20px; background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); border-radius: 10px; margin-bottom: 40px;'>
            <p style='color: #2c3e50; font-size: 1.1rem; line-height: 1.8; margin: 0;'>
                🌟 Chúng tôi là những Agribanker - những người giữ ánh sáng của niềm tin, lan tỏa tinh thần chuyển đổi số trên mọi miền đất nước
            </p>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("### 🌟 Thành viên nhóm")
    st.markdown("<br>", unsafe_allow_html=True)

    # Profile thành viên 1: Trần Ngọc Trúc Huỳnh
    col1, col2 = st.columns([1, 2])

    with col1:
        try:
            st.image("Tran Ngoc Truc Huynh.jpg", use_container_width=True)
        except:
            st.info("📸 Tran Ngoc Truc Huynh.jpg")

    with col2:
        st.markdown("""
            <div style='background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%); padding: 25px; border-radius: 15px; box-shadow: 0 8px 20px rgba(252, 182, 159, 0.3); height: 100%;'>
                <h3 style='color: #d63447; margin-top: 0; font-size: 1.8rem; border-bottom: 3px solid #d63447; padding-bottom: 10px;'>
                    🎯 Trần Ngọc Trúc Huỳnh
                </h3>
                <p style='color: #2c3e50; margin: 15px 0; font-size: 1.05rem;'>
                    <strong>🏢 Chức vụ:</strong> Giao dịch viên<br>
                    <strong>📍 Đơn vị:</strong> Agribank chi nhánh Tiền Giang
                </p>
                <div style='background: rgba(255, 255, 255, 0.6); padding: 15px; border-radius: 10px; margin-top: 15px;'>
                    <p style='color: #d63447; font-weight: 700; margin-bottom: 10px; font-size: 1.1rem;'>💼 Vai trò trong nhóm:</p>
                    <ul style='color: #2c3e50; margin: 0; padding-left: 20px; line-height: 1.8;'>
                        <li>Ý tưởng nâng cấp chương trình phiên bản 2.0</li>
                        <li>Kỹ thuật chính – Coder chính cho mô hình 2.0</li>
                        <li>Trailer giới thiệu mô hình nâng cấp</li>
                        <li>Phân chia, tổ chức công việc nhóm</li>
                        <li>Hỗ trợ kỹ thuật cho Version 1.0</li>
                        <li>Kịch bản & Thuyết trình Demo Version 1.0</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Profile thành viên 2: Nguyễn Hồng Cường
    col1, col2 = st.columns([1, 2])

    with col1:
        try:
            st.image("NGUYEN HONG CUONG.jpg", use_container_width=True)
        except:
            st.info("📸 NGUYEN HONG CUONG.jpg")

    with col2:
        st.markdown("""
            <div style='background: linear-gradient(135deg, #a1c4fd 0%, #c2e9fb 100%); padding: 25px; border-radius: 15px; box-shadow: 0 8px 20px rgba(161, 196, 253, 0.3); height: 100%;'>
                <h3 style='color: #2c5aa0; margin-top: 0; font-size: 1.8rem; border-bottom: 3px solid #2c5aa0; padding-bottom: 10px;'>
                    🎯 Nguyễn Hồng Cường
                </h3>
                <p style='color: #2c3e50; margin: 15px 0; font-size: 1.05rem;'>
                    <strong>🏢 Chức vụ:</strong> Trưởng phòng Kiểm tra – Kiểm soát Nội bộ<br>
                    <strong>📍 Đơn vị:</strong> Agribank chi nhánh Đông Hải Phòng
                </p>
                <div style='background: rgba(255, 255, 255, 0.6); padding: 15px; border-radius: 10px; margin-top: 15px;'>
                    <p style='color: #2c5aa0; font-weight: 700; margin-bottom: 10px; font-size: 1.1rem;'>💼 Vai trò trong nhóm:</p>
                    <ul style='color: #2c3e50; margin: 0; padding-left: 20px; line-height: 1.8;'>
                        <li>Kỹ thuật chính – Coder chính mô hình Version 1.0</li>
                        <li>Demo trực tiếp mô hình Version 1.0 trên sân khấu</li>
                        <li>Hỗ trợ kỹ thuật cho mô hình nâng cấp Version 2.0</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Profile thành viên 3: Nguyễn Trung Thành
    col1, col2 = st.columns([1, 2])

    with col1:
        try:
            st.image("NGUYEN TRUNG THANH.jpg", use_container_width=True)
        except:
            st.info("📸 NGUYEN TRUNG THANH.jpg")

    with col2:
        st.markdown("""
            <div style='background: linear-gradient(135deg, #ffeaa7 0%, #fdcb6e 100%); padding: 25px; border-radius: 15px; box-shadow: 0 8px 20px rgba(253, 203, 110, 0.3); height: 100%;'>
                <h3 style='color: #e17055; margin-top: 0; font-size: 1.8rem; border-bottom: 3px solid #e17055; padding-bottom: 10px;'>
                    🎯 Nguyễn Trung Thành
                </h3>
                <p style='color: #2c3e50; margin: 15px 0; font-size: 1.05rem;'>
                    <strong>🏢 Chức vụ:</strong> Phó trưởng Phòng Kế toán Ngân quỹ<br>
                    <strong>📍 Đơn vị:</strong> Agribank chi nhánh Hải Dương
                </p>
                <div style='background: rgba(255, 255, 255, 0.6); padding: 15px; border-radius: 10px; margin-top: 15px;'>
                    <p style='color: #e17055; font-weight: 700; margin-bottom: 10px; font-size: 1.1rem;'>💼 Vai trò trong nhóm:</p>
                    <ul style='color: #2c3e50; margin: 0; padding-left: 20px; line-height: 1.8;'>
                        <li>Hỗ trợ kỹ thuật cho mô hình Version 1.0</li>
                        <li>Thuyết trình sân khấu Demo Version 1.0</li>
                        <li>Thiết kế Poster mô hình Version 1.0</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Thông điệp kết thúc
    st.markdown("""
        <div style='text-align: center; padding: 30px; background: linear-gradient(135deg, #ff6b9d 0%, #c06c84 100%); border-radius: 15px; margin-top: 40px; box-shadow: 0 10px 30px rgba(255, 107, 157, 0.3);'>
            <h3 style='color: white; margin: 0 0 15px 0; font-size: 1.8rem;'>🚀 Sứ mệnh của chúng tôi</h3>
            <p style='color: #fff; font-size: 1.1rem; line-height: 1.8; margin: 0;'>
                Ứng dụng trí tuệ nhân tạo và công nghệ số để nâng cao hiệu quả hoạt động, quản trị rủi ro và chất lượng phục vụ khách hàng, góp phần hiện thực hóa chiến lược chuyển đổi số của Agribank.
            </p>
            <div style='margin-top: 20px; font-size: 2rem;'>
                💡 🎯 🌟 💼 🏆
            </div>
        </div>
    """, unsafe_allow_html=True)

# ========================================
# PREMIUM BANKING FOOTER
# ========================================
st.markdown("---")
footer_col1, footer_col2, footer_col3 = st.columns([2, 2, 1])

with footer_col1:
    st.markdown("""
    <div style='padding: 15px; text-align: left;'>
        <h4 style='color: #ff6b9d; margin-bottom: 10px;'>🏦 Chương Trình Đánh Giá Rủi Ro Tín Dụng</h4>
        <p style='color: #6b7280; font-size: 0.9rem; margin: 5px 0;'>
            Giải pháp AI tiên tiến cho phân tích tài chính doanh nghiệp
        </p>
        <p style='color: #6b7280; font-size: 0.85rem; margin: 5px 0;'>
            Authored by <strong>ÁNH SÁNG SỐ Team</strong> 
        </p>
    </div>
    """, unsafe_allow_html=True)

with footer_col2:
    st.markdown("""
    <div style='padding: 15px; text-align: left;'>
        <h4 style='color: #ff6b9d; margin-bottom: 10px;'>📊 Tính Năng Chính</h4>
        <ul style='color: #6b7280; font-size: 0.85rem; margin: 5px 0; padding-left: 20px;'>
            <li>Phân tích 14 chỉ số tài chính tự động</li>
            <li>Dự báo xác suất vỡ nợ (PD) và Phân tích chuyên sâu</li>
            <li>DashBoard Tài Chính Doanh Nghiệp tổng quan</li>
            <li>Tin tức tài chính cập nhật Real-Time</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

with footer_col3:
    st.markdown(f"""
    <div style='padding: 15px; text-align: center;'>
        <div style='font-size: 3rem; margin-bottom: 10px;'>💖</div>
        <p style='color: #ffb3c6; font-weight: 700; font-size: 0.9rem; margin: 5px 0;'>
            SWEET ANALYTICS
        </p>
        <p style='color: #6b7280; font-size: 0.75rem;'>
            Version 2.0 Premium
        </p>
    </div>
    """, unsafe_allow_html=True)

st.markdown(f"""
<div style='text-align: center; padding: 20px; margin-top: 20px;
            background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 100%);
            border-radius: 15px; box-shadow: 0 4px 15px rgba(255, 107, 157, 0.2);'>
    <p style='color: #ffffff; margin: 5px 0; font-size: 0.9rem; font-weight: 600;'>
        © {datetime.now().year} Credit Risk Assessment System | Developed with ❤️ using Streamlit
    </p>
    <p style='color: #fff0f5; margin: 5px 0; font-size: 0.85rem;'>
        🔒 Secure • 🚀 Fast • 🎯 Accurate • ✨ AI-Powered
    </p>
</div>
""", unsafe_allow_html=True)
